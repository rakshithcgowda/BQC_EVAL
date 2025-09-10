import sys
import datetime
import os
import sqlite3
import hashlib
import json
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from io import BytesIO
import logging
from typing import Dict, List, Tuple, Optional
from PyQt5.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout, 
                            QLabel, QLineEdit, QTextEdit, QComboBox, QSpinBox, QDoubleSpinBox, 
                            QCheckBox, QDateEdit, QPushButton, QGroupBox, QFormLayout, 
                            QScrollArea, QMessageBox, QFileDialog, QTabWidget, QRadioButton,
                            QButtonGroup, QFrame, QSizePolicy, QDialog, QDialogButtonBox,
                            QListWidget, QListWidgetItem, QDialogButtonBox, QProgressDialog)
from PyQt5.QtCore import Qt, QDate, pyqtSignal, QObject
from PyQt5.QtGui import QFont, QIcon, QPalette, QColor

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler("bqc_generator.log"),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# Get application data directory for portable database storage
def get_app_data_dir():
    """Get the application data directory in a portable way"""
    if os.name == 'nt':  # Windows
        app_data = os.getenv('LOCALAPPDATA')
        if not app_data:
            app_data = os.path.expanduser('~\\AppData\\Local')
    elif sys.platform == 'darwin':  # macOS
        app_data = os.path.expanduser('~/Library/Application Support')
    else:  # Linux and other Unix-like
        app_data = os.getenv('XDG_DATA_HOME', os.path.expanduser('~/.local/share'))
    
    # Create our app directory if it doesn't exist
    app_dir = os.path.join(app_data, 'BQCGenerator')
    if not os.path.exists(app_dir):
        os.makedirs(app_dir)
    
    return app_dir

# Constants
APP_DATA_DIR = get_app_data_dir()
DB_PATH = os.path.join(APP_DATA_DIR, "user_data.db")
LOG_PATH = os.path.join(APP_DATA_DIR, "bqc_generator.log")

# Update logging to use the new path
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler(LOG_PATH),
        logging.StreamHandler()
    ]
)

EMD_THRESHOLDS = [
    (50, 0),
    (100, 1),
    (500, 2.5),
    (1000, 5),
    (1500, 7.5),
    (2500, 10),
    (float('inf'), 20)
]

GROUP_OPTIONS = {
    '1': 'LPG',
    '2': 'GAS/HRS/CBG',
    '3': 'E&P GOODS',
    '4': 'E&P SERVICES',
    '6': 'LUBES',
    '7': 'PIPELINES',
    '8': 'BIOFUELS/DISPOSELS',
    '9': 'RETAIL/IS',
    '10': 'TRANSPORT'
}

TENDER_TYPES = ["Goods", "Service", "Works"]
MANUFACTURER_TYPES = [
    "Original Equipment Manufacturer", 
    "Authorized Channel Partner", 
    "Authorized Agent", 
    "Dealer", 
    "Authorized Distributor"
]

DIVISIBILITY_OPTIONS = ["Non-Divisible", "Divisible"]
DIVISION_PATTERNS = ["80:20", "70:20:10"]
PLATFORM_OPTIONS = ["GeM", "E-procurement"]

# Database setup
def setup_database():
    """Create database and tables if they don't exist"""
    try:
        conn = sqlite3.connect(DB_PATH)
        cursor = conn.cursor()
        
        # Create users table if it doesn't exist
        cursor.execute('''
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            username TEXT UNIQUE NOT NULL,
            password TEXT NOT NULL,
            email TEXT,
            full_name TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
        ''')
        
        # Check if bqc_data table exists
        cursor.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='bqc_data'")
        table_exists = cursor.fetchone()
        
        if not table_exists:
            # Create bqc_data table with all required columns
            cursor.execute('''
            CREATE TABLE bqc_data (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id INTEGER NOT NULL,
                ref_number TEXT NOT NULL,
                group_name TEXT,
                item_name TEXT,
                project_name TEXT,
                tender_description TEXT,
                pr_reference TEXT,
                tender_type TEXT,
                cec_estimate_incl_gst REAL,
                cec_date DATE,
                cec_estimate_excl_gst REAL,
                budget_details TEXT,
                tender_platform TEXT,
                scope_of_work TEXT,
                contract_period_years REAL,
                delivery_period TEXT,
                warranty_period TEXT,
                amc_period TEXT,
                payment_terms TEXT,
                manufacturer_types TEXT,
                supplying_capacity INTEGER,
                mse_relaxation INTEGER,
                similar_work_definition TEXT,
                annualized_value REAL,
                escalation_clause TEXT,
                divisibility TEXT,
                performance_security INTEGER,
                proposed_by TEXT,
                recommended_by TEXT,
                concurred_by TEXT,
                approved_by TEXT,
                amc_value REAL,
                has_amc INTEGER,
                correction_factor REAL,
                o_m_value REAL,
                o_m_period TEXT,
                has_om INTEGER,
                additional_details TEXT,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (user_id) REFERENCES users (id)
            )
            ''')
        else:
            # Get current table schema
            cursor.execute("PRAGMA table_info(bqc_data)")
            columns = [column[1] for column in cursor.fetchall()]
            
            # Add missing columns if they don't exist
            required_columns = [
                ('user_id', 'INTEGER NOT NULL'),
                ('ref_number', 'TEXT NOT NULL'),
                ('group_name', 'TEXT'),
                ('item_name', 'TEXT'),
                ('project_name', 'TEXT'),
                ('tender_description', 'TEXT'),
                ('pr_reference', 'TEXT'),
                ('tender_type', 'TEXT'),
                ('cec_estimate_incl_gst', 'REAL'),
                ('cec_date', 'DATE'),
                ('cec_estimate_excl_gst', 'REAL'),
                ('budget_details', 'TEXT'),
                ('tender_platform', 'TEXT'),
                ('scope_of_work', 'TEXT'),
                ('contract_period_years', 'REAL'),
                ('delivery_period', 'TEXT'),
                ('warranty_period', 'TEXT'),
                ('amc_period', 'TEXT'),
                ('payment_terms', 'TEXT'),
                ('manufacturer_types', 'TEXT'),
                ('supplying_capacity', 'INTEGER'),
                ('mse_relaxation', 'INTEGER'),
                ('similar_work_definition', 'TEXT'),
                ('annualized_value', 'REAL'),
                ('escalation_clause', 'TEXT'),
                ('divisibility', 'TEXT'),
                ('performance_security', 'INTEGER'),
                ('proposed_by', 'TEXT'),
                ('recommended_by', 'TEXT'),
                ('concurred_by', 'TEXT'),
                ('approved_by', 'TEXT'),
                ('amc_value', 'REAL'),
                ('has_amc', 'INTEGER'),
                ('correction_factor', 'REAL'),
                ('o_m_value', 'REAL'),
                ('o_m_period', 'TEXT'),
                ('has_om', 'INTEGER'),
                ('additional_details', 'TEXT'),
                ('created_at', 'TIMESTAMP DEFAULT CURRENT_TIMESTAMP'),
                ('updated_at', 'TIMESTAMP DEFAULT CURRENT_TIMESTAMP')
            ]
            
            # Add any missing columns
            for col_name, col_type in required_columns:
                if col_name not in columns:
                    try:
                        cursor.execute(f"ALTER TABLE bqc_data ADD COLUMN {col_name} {col_type}")
                        logger.info(f"Added missing column: {col_name}")
                    except sqlite3.OperationalError as e:
                        logger.error(f"Error adding column {col_name}: {str(e)}")
            
            # Handle the case where contract_period_months might exist instead of contract_period_years
            if 'contract_period_months' in columns and 'contract_period_years' not in columns:
                try:
                    cursor.execute("ALTER TABLE bqc_data RENAME COLUMN contract_period_months TO contract_period_years_old")
                    logger.info("Renamed contract_period_months to contract_period_years_old")
                except sqlite3.OperationalError as e:
                    logger.error(f"Error renaming column: {str(e)}")
            
            if 'contract_period_years' not in columns and 'contract_period_years_old' not in columns:
                try:
                    cursor.execute("ALTER TABLE bqc_data ADD COLUMN contract_period_years REAL DEFAULT 1")
                    logger.info("Added contract_period_years column")
                except sqlite3.OperationalError as e:
                    logger.error(f"Error adding contract_period_years: {str(e)}")
        
        conn.commit()
        conn.close()
        logger.info(f"Database setup completed successfully at {DB_PATH}")
        return True
    except Exception as e:
        logger.error(f"Error setting up database: {str(e)}")
        return False

def hash_password(password):
    """Hash a password using SHA-256 and hex"""
    return hashlib.sha256(password.encode()).hexdigest()

def sanitize_value(value):
    """Convert None values to 'NA' for database storage"""
    if value is None:
        return "NA"
    if isinstance(value, str) and value.strip() == "":
        return "NA"
    return value

class LoginDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Login")
        self.setMinimumWidth(300)
        self.user_id = None
        
        layout = QVBoxLayout()
        
        # Username
        username_layout = QHBoxLayout()
        username_label = QLabel("Username:")
        self.username_input = QLineEdit()
        username_layout.addWidget(username_label)
        username_layout.addWidget(self.username_input)
        layout.addLayout(username_layout)
        
        # Password
        password_layout = QHBoxLayout()
        password_label = QLabel("Password:")
        self.password_input = QLineEdit()
        self.password_input.setEchoMode(QLineEdit.Password)
        password_layout.addWidget(password_label)
        password_layout.addWidget(self.password_input)
        layout.addLayout(password_layout)
        
        # Buttons
        button_box = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        button_box.accepted.connect(self.accept)
        button_box.rejected.connect(self.reject)
        layout.addWidget(button_box)
        
        # Register button
        register_button = QPushButton("Register New User")
        register_button.clicked.connect(self.register_user)
        layout.addWidget(register_button)
        
        self.setLayout(layout)
    
    def register_user(self):
        """Open registration dialog"""
        dialog = RegistrationDialog(self)
        if dialog.exec_():
            QMessageBox.information(self, "Success", "Registration successful! Please login.")
    
    def accept(self):
        """Validate login credentials"""
        username = self.username_input.text().strip()
        password = self.password_input.text()
        
        if not username or not password:
            QMessageBox.warning(self, "Error", "Please enter both username and password")
            return
        
        try:
            conn = sqlite3.connect(DB_PATH)
            cursor = conn.cursor()
            
            # Check if user exists
            cursor.execute("SELECT id, password FROM users WHERE username = ?", (username,))
            result = cursor.fetchone()
            
            if result:
                user_id, stored_password = result
                hashed_password = hash_password(password)
                
                if stored_password == hashed_password:
                    self.user_id = user_id
                    super().accept()
                else:
                    QMessageBox.warning(self, "Error", "Invalid password")
            else:
                QMessageBox.warning(self, "Error", "User not found")
            
            conn.close()
        except Exception as e:
            logger.error(f"Login error: {str(e)}")
            QMessageBox.critical(self, "Error", "Database error occurred")

class RegistrationDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Register")
        self.setMinimumWidth(300)
        
        layout = QFormLayout()
        
        # Username
        self.username_input = QLineEdit()
        layout.addRow("Username:", self.username_input)
        
        # Password
        self.password_input = QLineEdit()
        self.password_input.setEchoMode(QLineEdit.Password)
        layout.addRow("Password:", self.password_input)
        
        # Confirm Password
        self.confirm_password_input = QLineEdit()
        self.confirm_password_input.setEchoMode(QLineEdit.Password)
        layout.addRow("Confirm Password:", self.confirm_password_input)
        
        # Email
        self.email_input = QLineEdit()
        layout.addRow("Email:", self.email_input)
        
        # Full Name
        self.full_name_input = QLineEdit()
        layout.addRow("Full Name:", self.full_name_input)
        
        # Buttons
        button_box = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        button_box.accepted.connect(self.accept)
        button_box.rejected.connect(self.reject)
        layout.addRow(button_box)
        
        self.setLayout(layout)
    
    def accept(self):
        """Validate and register new user"""
        username = self.username_input.text().strip()
        password = self.password_input.text()
        confirm_password = self.confirm_password_input.text()
        email = self.email_input.text().strip()
        full_name = self.full_name_input.text().strip()
        
        # Validation
        if not username or not password or not email or not full_name:
            QMessageBox.warning(self, "Error", "All fields are required")
            return
        
        if password != confirm_password:
            QMessageBox.warning(self, "Error", "Passwords do not match")
            return
        
        if len(password) < 6:
            QMessageBox.warning(self, "Error", "Password must be at least 6 characters")
            return
        
        try:
            conn = sqlite3.connect(DB_PATH)
            cursor = conn.cursor()
            
            # Check if username already exists
            cursor.execute("SELECT id FROM users WHERE username = ?", (username,))
            if cursor.fetchone():
                QMessageBox.warning(self, "Error", "Username already exists")
                conn.close()
                return
            
            # Insert new user
            hashed_password = hash_password(password)
            cursor.execute(
                "INSERT INTO users (username, password, email, full_name) VALUES (?, ?, ?, ?)",
                (username, hashed_password, email, full_name)
            )
            
            conn.commit()
            conn.close()
            
            super().accept()
        except Exception as e:
            logger.error(f"Registration error: {str(e)}")
            QMessageBox.critical(self, "Error", "Database error occurred")

class LoadDataDialog(QDialog):
    def __init__(self, user_id, parent=None):
        super().__init__(parent)
        self.user_id = user_id
        self.selected_id = None
        self.setWindowTitle("Load Saved Data")
        self.setMinimumWidth(500)
        self.setMinimumHeight(400)
        
        layout = QVBoxLayout()
        
        # List widget to show saved entries
        self.list_widget = QListWidget()
        self.list_widget.itemDoubleClicked.connect(self.accept)
        layout.addWidget(self.list_widget)
        
        # Buttons
        button_layout = QHBoxLayout()
        
        load_button = QPushButton("Load Selected")
        load_button.clicked.connect(self.accept)
        button_layout.addWidget(load_button)
        
        cancel_button = QPushButton("Cancel")
        cancel_button.clicked.connect(self.reject)
        button_layout.addWidget(cancel_button)
        
        layout.addLayout(button_layout)
        
        self.setLayout(layout)
        
        # Load saved data
        self.load_saved_data()
    
    def load_saved_data(self):
        """Load saved BQC data for the current user"""
        try:
            conn = sqlite3.connect(DB_PATH)
            cursor = conn.cursor()
            
            cursor.execute('''
            SELECT id, ref_number, tender_description, created_at 
            FROM bqc_data 
            WHERE user_id = ? 
            ORDER BY created_at DESC
            ''', (self.user_id,))
            
            results = cursor.fetchall()
            
            for row in results:
                id, ref_number, tender_description, created_at = row
                created_str = datetime.datetime.strptime(created_at, "%Y-%m-%d %H:%M:%S").strftime("%d/%m/%Y %H:%M")
                item_text = f"{ref_number} - {tender_description} ({created_str})"
                item = QListWidgetItem(item_text)
                item.setData(Qt.UserRole, id)
                self.list_widget.addItem(item)
            
            conn.close()
        except Exception as e:
            logger.error(f"Error loading saved data: {str(e)}")
            QMessageBox.critical(self, "Error", "Failed to load saved data")
    
    def accept(self):
        """Get the selected item ID"""
        selected_items = self.list_widget.selectedItems()
        if selected_items:
            self.selected_id = selected_items[0].data(Qt.UserRole)
            super().accept()
        else:
            QMessageBox.warning(self, "Error", "Please select an entry to load")

class BQCGeneratorApp(QMainWindow):
    def __init__(self, user_id):
        super().__init__()
        self.user_id = user_id
        self.setWindowTitle("BQC Document Generator")
        self.setGeometry(100, 100, 1200, 800)
        
        # Set application style
        self.setStyleSheet("""
            QMainWindow {
                background-color: #f5f5f5;
            }
            QTabWidget::pane {
                border: 1px solid #c5c5c5;
                background: white;
                border-radius: 4px;
            }
            QTabWidget::tab-bar {
                left: 5px;
            }
            QTabBar::tab {
                background: #e0e0e0;
                border: 1px solid #c5c5c5;
                border-bottom-color: #c2c7cb;
                border-top-left-radius: 4px;
                border-top-right-radius: 4px;
                min-width: 8ex;
                padding: 8px 12px;
                font-weight: bold;
            }
            QTabBar::tab:selected, QTabBar::tab:hover {
                background: #ffffff;
            }
            QTabBar::tab:selected {
                border-color: #3daee9;
                border-bottom-color: #ffffff;
            }
            QGroupBox {
                font-weight: bold;
                border: 1px solid #cccccc;
                border-radius: 5px;
                margin-top: 10px;
                padding-top: 10px;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 10px;
                padding: 0 5px 0 5px;
            }
            QPushButton {
                background-color: #3498db;
                color: white;
                border: none;
                padding: 8px 16px;
                border-radius: 4px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #2980b9;
            }
            QPushButton#generateButton {
                background-color: #2ecc71;
            }
            QPushButton#generateButton:hover {
                background-color: #27ae60;
            }
            QPushButton#clearButton {
                background-color: #e74c3c;
            }
            QPushButton#clearButton:hover {
                background-color: #c0392b;
            }
            QPushButton#saveButton {
                background-color: #f39c12;
            }
            QPushButton#saveButton:hover {
                background-color: #d68910;
            }
            QPushButton#loadButton {
                background-color: #9b59b6;
            }
            QPushButton#loadButton:hover {
                background-color: #8e44ad;
            }
            QLineEdit, QTextEdit, QComboBox, QSpinBox, QDoubleSpinBox, QDateEdit {
                border: 1px solid #cccccc;
                border-radius: 4px;
                padding: 5px;
                background-color: white;
            }
            QLabel {
                color: #333333;
            }
            QCheckBox {
                spacing: 8px;
            }
        """)
        
        # Initialize data dictionary
        self.data = {
            'ref_number': '',
            'group_name': '1 - Materials',
            'item_name': '',
            'project_name': '',
            'tender_description': '',
            'pr_reference': '',
            'tender_type': 'Goods',
            'cec_estimate_incl_gst': 0,
            'cec_date': datetime.date.today(),
            'cec_estimate_excl_gst': 0,
            'budget_details': '',
            'tender_platform': 'GeM',
            'scope_of_work': '',
            'contract_period_years': 1,
            'delivery_period': '',
            'warranty_period': '',
            'amc_period': '',
            'payment_terms': '',
            'manufacturer_types': ['Original Equipment Manufacturer'],
            'supplying_capacity': 30,
            'mse_relaxation': False,
            'similar_work_definition': '',
            'annualized_value': 0,
            'escalation_clause': '',
            'divisibility': 'Non-Divisible',
            'performance_security': 5,
            'proposed_by': 'XXXXX',
            'recommended_by': 'XXXXX',
            'concurred_by': 'Rajesh J.',
            'approved_by': 'Kani Amudhan N.',
            'amc_value': 0,
            'has_amc': False,
            'correction_factor': 0,
            'o_m_value': 0,
            'o_m_period': '',
            'has_om': False,
            'additional_details': ''
        }
        
        self.initUI()
        
    def initUI(self):
        # Create central widget and layout
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        main_layout = QVBoxLayout(central_widget)
        main_layout.setContentsMargins(20, 20, 20, 20)
        main_layout.setSpacing(15)
        
        # Create title
        title_label = QLabel("BQC Document Generator")
        title_font = QFont("Arial", 18, QFont.Bold)
        title_label.setFont(title_font)
        title_label.setAlignment(Qt.AlignCenter)
        title_label.setStyleSheet("color: #2c3e50; margin-bottom: 10px;")
        main_layout.addWidget(title_label)
        
        # Create subtitle
        subtitle_label = QLabel("Bid Qualification Criteria Generator for Procurement")
        subtitle_font = QFont("Arial", 10)
        subtitle_label.setFont(subtitle_font)
        subtitle_label.setAlignment(Qt.AlignCenter)
        subtitle_label.setStyleSheet("color: #7f8c8d; margin-bottom: 20px;")
        main_layout.addWidget(subtitle_label)
        
        # Create tab widget for different sections
        self.tab_widget = QTabWidget()
        self.tab_widget.setTabPosition(QTabWidget.North)
        main_layout.addWidget(self.tab_widget)
        
        # Create tabs
        self.preamble_tab = QWidget()
        self.scope_tab = QWidget()
        self.bqc_tab = QWidget()
        self.other_tab = QWidget()
        self.approval_tab = QWidget()
        
        self.tab_widget.addTab(self.preamble_tab, "Preamble")
        self.tab_widget.addTab(self.scope_tab, "Scope of Work")
        self.tab_widget.addTab(self.bqc_tab, "BQC Criteria")
        self.tab_widget.addTab(self.other_tab, "Other Sections")
        self.tab_widget.addTab(self.approval_tab, "Approval")
        
        # Setup each tab
        self.setup_preamble_tab(self.preamble_tab)
        self.setup_scope_tab(self.scope_tab)
        self.setup_bqc_tab(self.bqc_tab)
        self.setup_other_tab(self.other_tab)
        self.setup_approval_tab(self.approval_tab)
        
        # Create buttons
        button_layout = QHBoxLayout()
        button_layout.setSpacing(15)
        
        save_button = QPushButton("Save Data")
        save_button.setObjectName("saveButton")
        save_button.setMinimumHeight(40)
        save_button.clicked.connect(self.save_data)
        button_layout.addWidget(save_button)
        
        load_button = QPushButton("Load Data")
        load_button.setObjectName("loadButton")
        load_button.setMinimumHeight(40)
        load_button.clicked.connect(self.load_data)
        button_layout.addWidget(load_button)
        
        clear_button = QPushButton("Clear Form")
        clear_button.setObjectName("clearButton")
        clear_button.setMinimumHeight(40)
        clear_button.clicked.connect(self.clear_form)
        button_layout.addWidget(clear_button)
        
        generate_button = QPushButton("Generate Document")
        generate_button.setObjectName("generateButton")
        generate_button.setMinimumHeight(40)
        generate_button.clicked.connect(self.generate_document)
        button_layout.addWidget(generate_button)
        
        main_layout.addLayout(button_layout)
        
        # Status bar
        self.statusBar().showMessage("Ready")
    
    def setup_preamble_tab(self, tab):
        layout = QFormLayout(tab)
        layout.setFieldGrowthPolicy(QFormLayout.ExpandingFieldsGrow)
        layout.setLabelAlignment(Qt.AlignLeft)
        layout.setFormAlignment(Qt.AlignLeft)
        layout.setSpacing(12)
        layout.setContentsMargins(20, 20, 20, 20)
        
        # Ref Number
        ref_layout = QHBoxLayout()
        ref_label = QLabel("Ref Number:")
        ref_label.setMinimumWidth(150)
        ref_label.setStyleSheet("font-weight: bold;")
        ref_layout.addWidget(ref_label)
        
        self.ref_number_input = QLineEdit(self.data['ref_number'])
        self.ref_number_input.setPlaceholderText("XXXXXX")
        self.ref_number_input.setToolTip("Enter the unique reference number for this BQC document")
        self.ref_number_input.textChanged.connect(lambda text: self.update_data('ref_number', text))
        ref_layout.addWidget(self.ref_number_input)
        layout.addRow(ref_layout)
        
        # Group dropdown
        group_layout = QHBoxLayout()
        group_label = QLabel("Procurement Group:")
        group_label.setMinimumWidth(150)
        group_label.setStyleSheet("font-weight: bold;")
        group_layout.addWidget(group_label)
        
        self.group_combo = QComboBox()
        for key, value in GROUP_OPTIONS.items():
            self.group_combo.addItem(f"{key} - {value}", key)
        self.group_combo.setToolTip("Select the procurement group for this tender")
        self.group_combo.currentTextChanged.connect(self.update_group_name)
        group_layout.addWidget(self.group_combo)
        layout.addRow(group_layout)
        
        # Tender Description
        desc_label = QLabel("Tender Description:")
        desc_label.setMinimumWidth(150)
        desc_label.setStyleSheet("font-weight: bold;")
        
        self.tender_description_input = QTextEdit(self.data['tender_description'])
        self.tender_description_input.setPlaceholderText("Provide a detailed description of the tender")
        self.tender_description_input.setToolTip("Provide a detailed description of the tender requirements")
        self.tender_description_input.setMinimumHeight(80)
        self.tender_description_input.textChanged.connect(lambda: self.update_data('tender_description', self.tender_description_input.toPlainText()))
        layout.addRow(desc_label, self.tender_description_input)
        
        # PR Reference
        pr_layout = QHBoxLayout()
        pr_label = QLabel("PR Reference:")
        pr_label.setMinimumWidth(150)
        pr_label.setStyleSheet("font-weight: bold;")
        pr_layout.addWidget(pr_label)
        
        self.pr_reference_input = QLineEdit(self.data['pr_reference'])
        self.pr_reference_input.setPlaceholderText("Enter PR or email reference")
        self.pr_reference_input.setToolTip("Enter the PR reference or email reference for this tender")
        self.pr_reference_input.textChanged.connect(lambda text: self.update_data('pr_reference', text))
        pr_layout.addWidget(self.pr_reference_input)
        layout.addRow(pr_layout)
        
        # Tender Type
        tender_layout = QHBoxLayout()
        tender_label = QLabel("Type of Tender:")
        tender_label.setMinimumWidth(150)
        tender_label.setStyleSheet("font-weight: bold;")
        tender_layout.addWidget(tender_label)
        
        self.tender_type_combo = QComboBox()
        self.tender_type_combo.addItems(TENDER_TYPES)
        self.tender_type_combo.setCurrentText(self.data['tender_type'])
        self.tender_type_combo.setToolTip("Select the type of tender (Goods, Service, or Works)")
        self.tender_type_combo.currentTextChanged.connect(self.update_tender_type)
        tender_layout.addWidget(self.tender_type_combo)
        layout.addRow(tender_layout)
        
        # CEC Estimate (incl. GST)
        cec_incl_layout = QHBoxLayout()
        cec_incl_label = QLabel("CEC Estimate (incl. GST):")
        cec_incl_label.setMinimumWidth(150)
        cec_incl_label.setStyleSheet("font-weight: bold;")
        cec_incl_layout.addWidget(cec_incl_label)
        
        self.cec_incl_gst_input = QDoubleSpinBox()
        self.cec_incl_gst_input.setRange(0, 1000000)
        self.cec_incl_gst_input.setValue(self.data['cec_estimate_incl_gst'])
        self.cec_incl_gst_input.setSuffix(" Lakh")
        self.cec_incl_gst_input.setToolTip("Enter the CEC estimate including GST in Lakhs")
        self.cec_incl_gst_input.valueChanged.connect(lambda value: self.update_data('cec_estimate_incl_gst', value))
        cec_incl_layout.addWidget(self.cec_incl_gst_input)
        layout.addRow(cec_incl_layout)
        
        # CEC Date
        date_layout = QHBoxLayout()
        date_label = QLabel("CEC Date:")
        date_label.setMinimumWidth(150)
        date_label.setStyleSheet("font-weight: bold;")
        date_layout.addWidget(date_label)
        
        self.cec_date_input = QDateEdit()
        self.cec_date_input.setDate(QDate.currentDate())
        self.cec_date_input.setCalendarPopup(True)
        self.cec_date_input.setDisplayFormat("dd/MM/yyyy")
        self.cec_date_input.setToolTip("Select the date of the CEC estimate")
        self.cec_date_input.dateChanged.connect(self.update_cec_date)
        date_layout.addWidget(self.cec_date_input)
        layout.addRow(date_layout)
        
        # CEC Estimate (excl. GST)
        cec_excl_layout = QHBoxLayout()
        cec_excl_label = QLabel("CEC Estimate (excl. GST):")
        cec_excl_label.setMinimumWidth(150)
        cec_excl_label.setStyleSheet("font-weight: bold;")
        cec_excl_layout.addWidget(cec_excl_label)
        
        self.cec_excl_gst_input = QDoubleSpinBox()
        self.cec_excl_gst_input.setRange(0, 1000000)
        self.cec_excl_gst_input.setValue(self.data['cec_estimate_excl_gst'])
        self.cec_excl_gst_input.setSuffix(" Lakh")
        self.cec_excl_gst_input.setToolTip("Enter the CEC estimate excluding GST in Lakhs")
        self.cec_excl_gst_input.valueChanged.connect(lambda value: self.update_data('cec_estimate_excl_gst', value))
        cec_excl_layout.addWidget(self.cec_excl_gst_input)
        layout.addRow(cec_excl_layout)
        
        # Budget Details
        budget_layout = QHBoxLayout()
        budget_label = QLabel("Budget Details:")
        budget_label.setMinimumWidth(150)
        budget_label.setStyleSheet("font-weight: bold;")
        budget_layout.addWidget(budget_label)
        
        self.budget_details_input = QLineEdit(self.data['budget_details'])
        self.budget_details_input.setPlaceholderText("WBS/Revex")
        self.budget_details_input.setToolTip("Enter the budget details (WBS/Revex)")
        self.budget_details_input.textChanged.connect(lambda text: self.update_data('budget_details', text))
        budget_layout.addWidget(self.budget_details_input)
        layout.addRow(budget_layout)
        
        # Tender Platform
        platform_layout = QHBoxLayout()
        platform_label = QLabel("Tender Platform:")
        platform_label.setMinimumWidth(150)
        platform_label.setStyleSheet("font-weight: bold;")
        platform_layout.addWidget(platform_label)
        
        self.platform_combo = QComboBox()
        self.platform_combo.addItems(PLATFORM_OPTIONS)
        self.platform_combo.setCurrentText(self.data['tender_platform'])
        self.platform_combo.setToolTip("Select the tender platform (GeM or E-procurement)")
        self.platform_combo.currentTextChanged.connect(lambda text: self.update_data('tender_platform', text))
        platform_layout.addWidget(self.platform_combo)
        layout.addRow(platform_layout)
    
    def setup_scope_tab(self, tab):
        layout = QFormLayout(tab)
        layout.setFieldGrowthPolicy(QFormLayout.ExpandingFieldsGrow)
        layout.setLabelAlignment(Qt.AlignLeft)
        layout.setFormAlignment(Qt.AlignLeft)
        layout.setSpacing(12)
        layout.setContentsMargins(20, 20, 20, 20)
        
        # Scope of Work
        scope_label = QLabel("Brief Scope of Work:")
        scope_label.setMinimumWidth(150)
        scope_label.setStyleSheet("font-weight: bold;")
        
        self.scope_of_work_input = QTextEdit(self.data['scope_of_work'])
        self.scope_of_work_input.setPlaceholderText("Provide a brief scope of work or supply items")
        self.scope_of_work_input.setToolTip("Provide a brief description of the scope of work or supply items")
        self.scope_of_work_input.setMinimumHeight(80)
        self.scope_of_work_input.textChanged.connect(lambda: self.update_data('scope_of_work', self.scope_of_work_input.toPlainText()))
        layout.addRow(scope_label, self.scope_of_work_input)
        
        # Contract Period (Years)
        contract_layout = QHBoxLayout()
        contract_label = QLabel("Contract Period (Years):")
        contract_label.setMinimumWidth(150)
        contract_label.setStyleSheet("font-weight: bold;")
        contract_layout.addWidget(contract_label)
        
        self.contract_period_years_input = QDoubleSpinBox()
        self.contract_period_years_input.setRange(0.1, 20)
        self.contract_period_years_input.setValue(self.data['contract_period_years'])
        self.contract_period_years_input.setSingleStep(0.5)
        self.contract_period_years_input.setSuffix(" years")
        self.contract_period_years_input.setToolTip("Enter the contract period in years")
        self.contract_period_years_input.valueChanged.connect(self.update_contract_period_years)
        contract_layout.addWidget(self.contract_period_years_input)
        layout.addRow(contract_layout)
        
        # Annualized Value (calculated)
        annual_layout = QHBoxLayout()
        annual_label = QLabel("Annualized Value:")
        annual_label.setMinimumWidth(150)
        annual_label.setStyleSheet("font-weight: bold;")
        annual_layout.addWidget(annual_label)
        
        self.annualized_value_label = QLabel("Rs. 0.00 Lacs")
        self.annualized_value_label.setStyleSheet("color: #2c3e50; font-weight: bold;")
        annual_layout.addWidget(self.annualized_value_label)
        layout.addRow(annual_layout)
        
        # Maintenance Options Group
        maintenance_group = QGroupBox("Maintenance Options")
        maintenance_layout = QVBoxLayout()
        maintenance_layout.setSpacing(12)
        maintenance_layout.setContentsMargins(15, 15, 15, 15)
        
        # AMC/CAMC Checkbox
        amc_check_layout = QHBoxLayout()
        self.has_amc_checkbox = QCheckBox("Has AMC/CAMC?")
        self.has_amc_checkbox.setChecked(self.data['has_amc'])
        self.has_amc_checkbox.setToolTip("Check if this tender includes Annual Maintenance Contract")
        self.has_amc_checkbox.stateChanged.connect(self.update_has_amc)
        amc_check_layout.addWidget(self.has_amc_checkbox)
        maintenance_layout.addLayout(amc_check_layout)
        
        # AMC Period (only visible when has_amc is checked)
        amc_period_layout = QHBoxLayout()
        amc_period_label = QLabel("AMC Period:")
        amc_period_label.setMinimumWidth(150)
        amc_period_label.setStyleSheet("font-weight: bold;")
        amc_period_layout.addWidget(amc_period_label)
        
        self.amc_period_input = QLineEdit(self.data['amc_period'])
        self.amc_period_input.setPlaceholderText("e.g., 3 years")
        self.amc_period_input.setToolTip("Enter the AMC/CAMC period in years")
        self.amc_period_input.textChanged.connect(lambda text: self.update_data('amc_period', text))
        amc_period_layout.addWidget(self.amc_period_input)
        maintenance_layout.addLayout(amc_period_layout)
        
        # AMC Value (only visible when has_amc is checked)
        amc_value_layout = QHBoxLayout()
        amc_value_label = QLabel("AMC Value:")
        amc_value_label.setMinimumWidth(150)
        amc_value_label.setStyleSheet("font-weight: bold;")
        amc_value_layout.addWidget(amc_value_label)
        
        self.amc_value_input = QDoubleSpinBox()
        self.amc_value_input.setRange(0, 1000000)
        self.amc_value_input.setValue(self.data['amc_value'])
        self.amc_value_input.setSuffix(" Lakh")
        self.amc_value_input.setToolTip("Enter the AMC value in Lakhs")
        self.amc_value_input.valueChanged.connect(lambda value: self.update_data('amc_value', value))
        amc_value_layout.addWidget(self.amc_value_input)
        maintenance_layout.addLayout(amc_value_layout)
        
        maintenance_group.setLayout(maintenance_layout)
        layout.addRow(maintenance_group)
        
        # Payment Terms
        payment_layout = QHBoxLayout()
        payment_label = QLabel("Payment Terms:")
        payment_label.setMinimumWidth(150)
        payment_label.setStyleSheet("font-weight: bold;")
        payment_layout.addWidget(payment_label)
        
        self.payment_terms_input = QLineEdit(self.data['payment_terms'])
        self.payment_terms_input.setPlaceholderText("e.g., Within 30 days")
        self.payment_terms_input.setToolTip("Enter payment terms if different from standard (within 30 days)")
        self.payment_terms_input.textChanged.connect(lambda text: self.update_data('payment_terms', text))
        payment_layout.addWidget(self.payment_terms_input)
        layout.addRow(payment_layout)
        
        # Goods-specific fields
        self.goods_group = QGroupBox("Goods-Specific Fields")
        goods_layout = QFormLayout()
        goods_layout.setFieldGrowthPolicy(QFormLayout.ExpandingFieldsGrow)
        goods_layout.setLabelAlignment(Qt.AlignLeft)
        goods_layout.setFormAlignment(Qt.AlignLeft)
        goods_layout.setSpacing(12)
        goods_layout.setContentsMargins(15, 15, 15, 15)
        
        # Delivery Period
        delivery_layout = QHBoxLayout()
        delivery_label = QLabel("Delivery Period:")
        delivery_label.setMinimumWidth(150)
        delivery_label.setStyleSheet("font-weight: bold;")
        delivery_layout.addWidget(delivery_label)
        
        self.delivery_period_input = QLineEdit(self.data['delivery_period'])
        self.delivery_period_input.setPlaceholderText("e.g., 30 days")
        self.delivery_period_input.setToolTip("Enter the delivery period for goods")
        self.delivery_period_input.textChanged.connect(lambda text: self.update_data('delivery_period', text))
        delivery_layout.addWidget(self.delivery_period_input)
        goods_layout.addRow(delivery_layout)
        
        # Warranty Period
        warranty_layout = QHBoxLayout()
        warranty_label = QLabel("Warranty Period:")
        warranty_label.setMinimumWidth(150)
        warranty_label.setStyleSheet("font-weight: bold;")
        warranty_layout.addWidget(warranty_label)
        
        self.warranty_period_input = QLineEdit(self.data['warranty_period'])
        self.warranty_period_input.setPlaceholderText("e.g., 12 months")
        self.warranty_period_input.setToolTip("Enter the warranty period for goods")
        self.warranty_period_input.textChanged.connect(lambda text: self.update_data('warranty_period', text))
        warranty_layout.addWidget(self.warranty_period_input)
        goods_layout.addRow(warranty_layout)
        
        self.goods_group.setLayout(goods_layout)
        layout.addRow(self.goods_group)
        
        # Initially set visibility
        self.update_amc_visibility()
        self.goods_group.setVisible(self.data['tender_type'] == 'Goods')
    
    def setup_bqc_tab(self, tab):
        layout = QVBoxLayout(tab)
        layout.setSpacing(15)
        layout.setContentsMargins(20, 20, 20, 20)
        
        # Technical Criteria
        tech_group = QGroupBox("Technical Criteria")
        tech_layout = QVBoxLayout()
        tech_layout.setSpacing(12)
        tech_layout.setContentsMargins(15, 15, 15, 15)
        
        # Goods-specific technical criteria
        self.goods_tech_group = QGroupBox("For Goods")
        goods_tech_layout = QVBoxLayout()
        goods_tech_layout.setSpacing(12)
        goods_tech_layout.setContentsMargins(15, 15, 15, 15)
        
        # Manufacturer Types
        manufacturer_layout = QHBoxLayout()
        manufacturer_layout.addWidget(QLabel("Manufacturer Types:"))
        manufacturer_layout.addSpacing(10)
        self.manufacturer_checkboxes = []
        for m_type in MANUFACTURER_TYPES:
            cb = QCheckBox(m_type)
            cb.setChecked(m_type in self.data['manufacturer_types'])
            cb.setToolTip(f"Select if bidders can be {m_type}")
            cb.stateChanged.connect(self.update_manufacturer_types)
            manufacturer_layout.addWidget(cb)
            self.manufacturer_checkboxes.append(cb)
        goods_tech_layout.addLayout(manufacturer_layout)
        
        # Supplying Capacity
        supplying_layout = QHBoxLayout()
        supplying_layout.addWidget(QLabel("Quantity to be supplied:"))
        supplying_layout.addSpacing(10)
        self.supplying_capacity_input = QSpinBox()
        self.supplying_capacity_input.setRange(0, 1000)
        self.supplying_capacity_input.setValue(self.data['supplying_capacity'])
        self.supplying_capacity_input.setSuffix("")
        self.supplying_capacity_input.setToolTip("Enter the base value for supplying capacity (30% of this will be used)")
        self.supplying_capacity_input.valueChanged.connect(self.update_supplying_capacity)
        supplying_layout.addWidget(self.supplying_capacity_input)
        goods_tech_layout.addLayout(supplying_layout)
        
        # Display calculated supplying capacity
        self.calculated_capacity_label = QLabel("Calculated Supplying Capacity: 30% of 30 = 9")
        self.calculated_capacity_label.setStyleSheet("color: #3498db; font-style: italic; padding: 5px;")
        goods_tech_layout.addWidget(self.calculated_capacity_label)
        
        # MSE Relaxation
        self.mse_relaxation_checkbox = QCheckBox("Apply MSE Relaxation (15%)")
        self.mse_relaxation_checkbox.setChecked(self.data['mse_relaxation'])
        self.mse_relaxation_checkbox.setToolTip("Check to apply 15% relaxation for MSE bidders")
        self.mse_relaxation_checkbox.stateChanged.connect(self.update_mse_relaxation)
        goods_tech_layout.addWidget(self.mse_relaxation_checkbox)
        
        # MSE Calculation Display (initially hidden)
        self.mse_calculation_label = QLabel("")
        self.mse_calculation_label.setStyleSheet("color: #3498db; font-style: italic; padding: 5px;")
        self.mse_calculation_label.setVisible(False)
        goods_tech_layout.addWidget(self.mse_calculation_label)
        
        self.goods_tech_group.setLayout(goods_tech_layout)
        tech_layout.addWidget(self.goods_tech_group)
        
        # Service/Works-specific technical criteria
        self.service_tech_group = QGroupBox("For Service/Works")
        service_tech_layout = QVBoxLayout()
        service_tech_layout.setSpacing(12)
        service_tech_layout.setContentsMargins(15, 15, 15, 15)
        
        # Similar Work Definition
        service_tech_layout.addWidget(QLabel("Definition of Similar Work:"))
        self.similar_work_input = QTextEdit(self.data['similar_work_definition'])
        self.similar_work_input.setPlaceholderText("Define what constitutes similar work for this tender")
        self.similar_work_input.setToolTip("Provide a clear definition of what constitutes similar work")
        self.similar_work_input.setMinimumHeight(80)
        self.similar_work_input.textChanged.connect(lambda: self.update_data('similar_work_definition', self.similar_work_input.toPlainText()))
        service_tech_layout.addWidget(self.similar_work_input)
        
        # Experience Requirements
        self.experience_req_label = QLabel("Experience Requirements:")
        self.experience_req_label.setStyleSheet("font-weight: bold; margin-top: 10px;")
        service_tech_layout.addWidget(self.experience_req_label)
        
        self.service_tech_group.setLayout(service_tech_layout)
        tech_layout.addWidget(self.service_tech_group)
        
        tech_group.setLayout(tech_layout)
        layout.addWidget(tech_group)
        
        # Financial Criteria
        financial_group = QGroupBox("Financial Criteria")
        financial_layout = QVBoxLayout()
        financial_layout.setSpacing(12)
        financial_layout.setContentsMargins(15, 15, 15, 15)
        
        # Turnover Requirement
        self.turnover_req_label = QLabel("Annual Turnover Requirement is: Rs. 0.00 Lacs")
        self.turnover_req_label.setStyleSheet("color: #2c3e50; font-weight: bold;")
        financial_layout.addWidget(self.turnover_req_label)
        
        financial_group.setLayout(financial_layout)
        layout.addWidget(financial_group)
        
        # Update visibility based on tender type
        self.update_bqc_tab_visibility()
    
    def setup_other_tab(self, tab):
        layout = QFormLayout(tab)
        layout.setFieldGrowthPolicy(QFormLayout.ExpandingFieldsGrow)
        layout.setLabelAlignment(Qt.AlignLeft)
        layout.setFormAlignment(Qt.AlignLeft)
        layout.setSpacing(12)
        layout.setContentsMargins(20, 20, 20, 20)
        
        # Escalation Clause
        escalation_label = QLabel("Escalation Clause:")
        escalation_label.setMinimumWidth(150)
        escalation_label.setStyleSheet("font-weight: bold;")
        
        self.escalation_clause_input = QTextEdit(self.data['escalation_clause'])
        self.escalation_clause_input.setPlaceholderText("Enter escalation/de-escalation clause details")
        self.escalation_clause_input.setToolTip("Enter details of escalation or de-escalation clause if applicable")
        self.escalation_clause_input.setMinimumHeight(80)
        self.escalation_clause_input.textChanged.connect(lambda: self.update_data('escalation_clause', self.escalation_clause_input.toPlainText()))
        layout.addRow(escalation_label, self.escalation_clause_input)
        
        # Additional Details
        additional_label = QLabel("Additional Details:")
        additional_label.setMinimumWidth(150)
        additional_label.setStyleSheet("font-weight: bold;")
        
        self.additional_details_input = QTextEdit(self.data['additional_details'])
        self.additional_details_input.setPlaceholderText("Enter any additional details or requirements")
        self.additional_details_input.setToolTip("Enter any additional details or special requirements for this tender")
        self.additional_details_input.setMinimumHeight(80)
        self.additional_details_input.textChanged.connect(lambda: self.update_data('additional_details', self.additional_details_input.toPlainText()))
        layout.addRow(additional_label, self.additional_details_input)
        
        # Divisibility
        div_layout = QHBoxLayout()
        div_label = QLabel("Divisibility:")
        div_label.setMinimumWidth(150)
        div_label.setStyleSheet("font-weight: bold;")
        div_layout.addWidget(div_label)
        
        self.divisibility_combo = QComboBox()
        self.divisibility_combo.addItems(DIVISIBILITY_OPTIONS)
        self.divisibility_combo.setCurrentText(self.data['divisibility'])
        self.divisibility_combo.setToolTip("Select if the tender is divisible or non-divisible")
        self.divisibility_combo.currentTextChanged.connect(self.update_divisibility)
        div_layout.addWidget(self.divisibility_combo)
        layout.addRow(div_layout)
        
        # Correction Factor (for divisibility) - MOVED HERE
        self.correction_factor_group = QGroupBox("Divisibility Settings")
        correction_layout = QFormLayout()
        correction_layout.setFieldGrowthPolicy(QFormLayout.ExpandingFieldsGrow)
        correction_layout.setLabelAlignment(Qt.AlignLeft)
        correction_layout.setFormAlignment(Qt.AlignLeft)
        correction_layout.setSpacing(12)
        correction_layout.setContentsMargins(15, 15, 15, 15)
        
        self.correction_factor_input = QDoubleSpinBox()
        self.correction_factor_input.setRange(0, 1)
        self.correction_factor_input.setSingleStep(0.01)
        self.correction_factor_input.setValue(self.data['correction_factor'])
        self.correction_factor_input.setToolTip("Enter the correction factor for divisible tenders")
        self.correction_factor_input.valueChanged.connect(lambda value: self.update_data('correction_factor', value))
        correction_layout.addRow("Correction Factor:", self.correction_factor_input)
        
        self.correction_factor_group.setLayout(correction_layout)
        layout.addRow(self.correction_factor_group)
        
        # EMD Calculation Preview
        emd_layout = QHBoxLayout()
        emd_label = QLabel("EMD Preview:")
        emd_label.setMinimumWidth(150)
        emd_label.setStyleSheet("font-weight: bold;")
        emd_layout.addWidget(emd_label)
        
        self.emd_preview_label = QLabel("EMD will be automatically calculated as: Rs. 0.00 Lacs")
        self.emd_preview_label.setStyleSheet("color: #2c3e50; font-weight: bold;")
        emd_layout.addWidget(self.emd_preview_label)
        layout.addRow(emd_layout)
        
        # Performance Security
        ps_layout = QHBoxLayout()
        ps_label = QLabel("Performance Security:")
        ps_label.setMinimumWidth(150)
        ps_label.setStyleSheet("font-weight: bold;")
        ps_layout.addWidget(ps_label)
        
        self.performance_security_input = QSpinBox()
        self.performance_security_input.setRange(0, 20)
        self.performance_security_input.setValue(self.data['performance_security'])
        self.performance_security_input.setSuffix("%")
        self.performance_security_input.setToolTip("Enter the performance security percentage")
        self.performance_security_input.valueChanged.connect(lambda value: self.update_data('performance_security', value))
        ps_layout.addWidget(self.performance_security_input)
        layout.addRow(ps_layout)
        
        # Initially set visibility
        self.correction_factor_group.setVisible(self.data['divisibility'] == 'Divisible')
    
    def setup_approval_tab(self, tab):
        layout = QFormLayout(tab)
        layout.setFieldGrowthPolicy(QFormLayout.ExpandingFieldsGrow)
        layout.setLabelAlignment(Qt.AlignLeft)
        layout.setFormAlignment(Qt.AlignLeft)
        layout.setSpacing(12)
        layout.setContentsMargins(20, 20, 20, 20)
        
        # Proposed By
        proposed_layout = QHBoxLayout()
        proposed_label = QLabel("Proposed By:")
        proposed_label.setMinimumWidth(150)
        proposed_label.setStyleSheet("font-weight: bold;")
        proposed_layout.addWidget(proposed_label)
        
        self.proposed_by_input = QLineEdit(self.data['proposed_by'])
        self.proposed_by_input.setToolTip("Enter the name of the person proposing this BQC")
        self.proposed_by_input.textChanged.connect(lambda text: self.update_data('proposed_by', text))
        proposed_layout.addWidget(self.proposed_by_input)
        layout.addRow(proposed_layout)
        
        # Recommended By
        recommended_layout = QHBoxLayout()
        recommended_label = QLabel("Recommended By:")
        recommended_label.setMinimumWidth(150)
        recommended_label.setStyleSheet("font-weight: bold;")
        recommended_layout.addWidget(recommended_label)
        
        self.recommended_by_input = QLineEdit(self.data['recommended_by'])
        self.recommended_by_input.setToolTip("Enter the name of the person recommending this BQC")
        self.recommended_by_input.textChanged.connect(lambda text: self.update_data('recommended_by', text))
        recommended_layout.addWidget(self.recommended_by_input)
        layout.addRow(recommended_layout)
        
        # Concurred By
        concurred_layout = QHBoxLayout()
        concurred_label = QLabel("Concurred By:")
        concurred_label.setMinimumWidth(150)
        concurred_label.setStyleSheet("font-weight: bold;")
        concurred_layout.addWidget(concurred_label)
        
        self.concurred_by_input = QLineEdit(self.data['concurred_by'])
        self.concurred_by_input.setToolTip("Enter the name of the person concurring this BQC")
        self.concurred_by_input.textChanged.connect(lambda text: self.update_data('concurred_by', text))
        concurred_layout.addWidget(self.concurred_by_input)
        layout.addRow(concurred_layout)
        
        # Approved By
        approved_layout = QHBoxLayout()
        approved_label = QLabel("Approved By:")
        approved_label.setMinimumWidth(150)
        approved_label.setStyleSheet("font-weight: bold;")
        approved_layout.addWidget(approved_label)
        
        self.approved_by_input = QLineEdit(self.data['approved_by'])
        self.approved_by_input.setToolTip("Enter the name of the person approving this BQC")
        self.approved_by_input.textChanged.connect(lambda text: self.update_data('approved_by', text))
        approved_layout.addWidget(self.approved_by_input)
        layout.addRow(approved_layout)
    
    def update_data(self, key, value):
        self.data[key] = value
        self.update_calculated_values()
    
    def update_group_name(self, text):
        key = text.split(" - ")[0]
        self.data['group_name'] = f"{key} - {GROUP_OPTIONS[key]}"
    
    def update_tender_type(self, text):
        self.data['tender_type'] = text
        
        # Update visibility of goods-specific fields
        self.goods_group.setVisible(text == 'Goods')
        
        # Update visibility in BQC tab
        self.update_bqc_tab_visibility()
        
        # Update performance security default
        if text in ['Goods', 'Services']:
            self.performance_security_input.setValue(5)
        else:
            self.performance_security_input.setValue(10)
    
    def update_cec_date(self, date):
        self.data['cec_date'] = date.toPyDate()
    
    def update_contract_period_years(self, value):
        self.data['contract_period_years'] = value
        self.update_calculated_values()
    
    def update_supplying_capacity(self, value):
        # Store the base value
        self.data['supplying_capacity'] = value
        
        # Calculate 30% of the base value
        calculated_value = int(value * 0.3)
        
        # Update the display label
        self.calculated_capacity_label.setText(f"Calculated Supplying Capacity: 30% of {value} = {calculated_value}")
        
        # Update MSE calculation if applicable
        if self.data['mse_relaxation']:
            self.update_mse_relaxation_display()
        
        self.update_calculated_values()
    
    def update_manufacturer_types(self):
        selected_types = []
        for i, cb in enumerate(self.manufacturer_checkboxes):
            if cb.isChecked():
                selected_types.append(MANUFACTURER_TYPES[i])
        self.data['manufacturer_types'] = selected_types
    
    def update_mse_relaxation(self, state):
        self.data['mse_relaxation'] = (state == Qt.Checked)
        self.update_mse_relaxation_display()
        self.update_calculated_values()
        
        # Show/hide MSE calculation label
        self.mse_calculation_label.setVisible(self.data['mse_relaxation'])
    
    def update_mse_relaxation_display(self):
        if self.data['mse_relaxation']:
            base_value = self.data['supplying_capacity']
            calculated_value = int(base_value * 0.3)
            relaxed_value = int(calculated_value * 0.85)  # 15% relaxation
            self.mse_calculation_label.setText(
                f"MSE Relaxation Calculation: {calculated_value} × (1 - 15%) = {relaxed_value}"
            )
        else:
            self.mse_calculation_label.setText("")
    
    def update_has_amc(self, state):
        self.data['has_amc'] = (state == Qt.Checked)
        self.update_amc_visibility()
    
    def update_amc_visibility(self):
        # Show/hide AMC fields based on checkbox
        has_amc = self.data['has_amc']
        
        # Set visibility for AMC period input and its label
        if hasattr(self, 'amc_period_input'):
            self.amc_period_input.setVisible(has_amc)
            # Find the label for this input and hide/show it too
            if hasattr(self, 'scope_tab') and self.scope_tab.layout():
                for i in range(self.scope_tab.layout().count()):
                    item = self.scope_tab.layout().itemAt(i)
                    if item and item.widget() and isinstance(item.widget(), QLabel):
                        if item.widget().text() == "AMC Period:":
                            item.widget().setVisible(has_amc)
                            break
        
        # Set visibility for AMC value input and its label
        if hasattr(self, 'amc_value_input'):
            self.amc_value_input.setVisible(has_amc)
            # Find the label for this input and hide/show it too
            if hasattr(self, 'scope_tab') and self.scope_tab.layout():
                for i in range(self.scope_tab.layout().count()):
                    item = self.scope_tab.layout().itemAt(i)
                    if item and item.widget() and isinstance(item.widget(), QLabel):
                        if item.widget().text() == "AMC Value:":
                            item.widget().setVisible(has_amc)
                            break
    
    def update_divisibility(self, text):
        self.data['divisibility'] = text
        # Update visibility of correction factor
        is_divisible = (text == 'Divisible')
        if hasattr(self, 'correction_factor_group'):
            self.correction_factor_group.setVisible(is_divisible)
        self.update_calculated_values()
    
    def update_bqc_tab_visibility(self):
        tender_type = self.data['tender_type']
        self.goods_tech_group.setVisible(tender_type == 'Goods')
        self.service_tech_group.setVisible(tender_type in ['Service', 'Works'])
    
    def update_calculated_values(self):
        # Update annualized value based on contract period in years
        if self.data['contract_period_years'] > 0:
            annualized_value = self.data['cec_estimate_excl_gst'] / self.data['contract_period_years']
            self.data['annualized_value'] = annualized_value
            self.annualized_value_label.setText(f"Rs. {annualized_value:.2f} Lacs")
        else:
            self.annualized_value_label.setText("Rs. 0.00 Lacs")
        
        # Update turnover requirement with refined logic
        base_percentage = 0.3
        
        # Apply correction factor when divisible for all tender types
        if self.data['divisibility'] == 'Divisible':
            base_percentage = 0.3 * (1 + self.data['correction_factor'])
        
        # Calculate turnover requirement considering AMC value
        maintenance_value = 0
        maintenance_text = ""
        
        if self.data['has_amc']:
            maintenance_value = self.data['amc_value']
            maintenance_text = "AMC"
        
        if maintenance_value > 0:
            # Calculate as base_percentage of (CEC_incl_gst - maintenance_value)
            turnover_requirement = base_percentage * (self.data['cec_estimate_incl_gst'] - maintenance_value)
            self.turnover_req_label.setText(
                f"Annual Turnover Requirement ({base_percentage*100:.0f}% of CEC-{maintenance_text}): Rs. {turnover_requirement:.2f} Lacs"
            )
        else:
            # Calculate as base_percentage of CEC
            turnover_requirement = base_percentage * self.data['cec_estimate_excl_gst']
            self.turnover_req_label.setText(
                f"Annual Turnover Requirement ({base_percentage*100:.0f}% of CEC): Rs. {turnover_requirement:.2f} Lacs"
            )
        
        # Update EMD preview
        emd_amount = calculate_emd(self.data['cec_estimate_excl_gst'], self.data['tender_type'])
        if emd_amount == 0:
            emd_text = "Nil"
        else:
            emd_text = f"{emd_amount} Lacs"
        self.emd_preview_label.setText(f"EMD will be automatically calculated as: Rs. {emd_text}")
        
        # Update experience requirements for Service/Works
        if self.data['tender_type'] in ['Service', 'Works']:
            # Apply correction factor if divisible
            if self.data['divisibility'] == 'Divisible':
                correction_factor = self.data['correction_factor']
                option_a_percent = 0.4 * (1 + correction_factor)
                option_b_percent = 0.5 * (1 + correction_factor)
                option_c_percent = 0.8 * (1 + correction_factor)
            else:
                option_a_percent = 0.4
                option_b_percent = 0.5
                option_c_percent = 0.8
            
            option_a_value = option_a_percent * self.data['cec_estimate_incl_gst']
            option_b_value = option_b_percent * self.data['cec_estimate_incl_gst']
            option_c_value = option_c_percent * self.data['cec_estimate_incl_gst']
            
            experience_text = (
                f"Option A ({option_a_percent*100:.0f}%): Rs. {option_a_value:.2f} Lacs\n"
                f"Option B ({option_b_percent*100:.0f}%): Rs. {option_b_value:.2f} Lacs\n"
                f"Option C ({option_c_percent*100:.0f}%): Rs. {option_c_value:.2f} Lacs"
            )
            self.experience_req_label.setText(experience_text)
    
    def save_data(self):
        """Save current form data to database"""
        if not self.data['ref_number']:
            QMessageBox.warning(self, "Error", "Reference Number is required to save data")
            return
        
        # Show progress dialog
        progress = QProgressDialog("Saving data...", "Cancel", 0, 0, self)
        progress.setWindowModality(Qt.WindowModal)
        progress.show()
        
        try:
            conn = sqlite3.connect(DB_PATH)
            cursor = conn.cursor()
            
            # Check if record with this ref_number already exists for this user
            cursor.execute('''
            SELECT id FROM bqc_data 
            WHERE user_id = ? AND ref_number = ?
            ''', (self.user_id, self.data['ref_number']))
            
            existing_record = cursor.fetchone()
            
            # Convert complex data types to strings
            manufacturer_types_json = json.dumps(self.data['manufacturer_types'])
            
            # Prepare data for database - convert None/empty to "NA"
            db_data = {
                'group_name': sanitize_value(self.data['group_name']),
                'item_name': sanitize_value(self.data['item_name']),
                'project_name': sanitize_value(self.data['project_name']),
                'tender_description': sanitize_value(self.data['tender_description']),
                'pr_reference': sanitize_value(self.data['pr_reference']),
                'tender_type': sanitize_value(self.data['tender_type']),
                'cec_estimate_incl_gst': self.data['cec_estimate_incl_gst'],
                'cec_date': self.data['cec_date'].isoformat(),
                'cec_estimate_excl_gst': self.data['cec_estimate_excl_gst'],
                'budget_details': sanitize_value(self.data['budget_details']),
                'tender_platform': sanitize_value(self.data['tender_platform']),
                'scope_of_work': sanitize_value(self.data['scope_of_work']),
                'contract_period_years': self.data['contract_period_years'],
                'delivery_period': sanitize_value(self.data['delivery_period']),
                'warranty_period': sanitize_value(self.data['warranty_period']),
                'amc_period': sanitize_value(self.data['amc_period']),
                'payment_terms': sanitize_value(self.data['payment_terms']),
                'manufacturer_types': manufacturer_types_json,
                'supplying_capacity': self.data['supplying_capacity'],
                'mse_relaxation': 1 if self.data['mse_relaxation'] else 0,
                'similar_work_definition': sanitize_value(self.data['similar_work_definition']),
                'annualized_value': self.data['annualized_value'],
                'escalation_clause': sanitize_value(self.data['escalation_clause']),
                'divisibility': sanitize_value(self.data['divisibility']),
                'performance_security': self.data['performance_security'],
                'proposed_by': sanitize_value(self.data['proposed_by']),
                'recommended_by': sanitize_value(self.data['recommended_by']),
                'concurred_by': sanitize_value(self.data['concurred_by']),
                'approved_by': sanitize_value(self.data['approved_by']),
                'amc_value': self.data['amc_value'],
                'has_amc': 1 if self.data['has_amc'] else 0,
                'correction_factor': self.data['correction_factor'],
                'o_m_value': self.data['o_m_value'],
                'o_m_period': sanitize_value(self.data['o_m_period']),
                'has_om': 1 if self.data['has_om'] else 0,
                'additional_details': sanitize_value(self.data['additional_details'])
            }
            
            if existing_record:
                # Update existing record
                cursor.execute('''
                UPDATE bqc_data SET
                    group_name = ?,
                    item_name = ?,
                    project_name = ?,
                    tender_description = ?,
                    pr_reference = ?,
                    tender_type = ?,
                    cec_estimate_incl_gst = ?,
                    cec_date = ?,
                    cec_estimate_excl_gst = ?,
                    budget_details = ?,
                    tender_platform = ?,
                    scope_of_work = ?,
                    contract_period_years = ?,
                    delivery_period = ?,
                    warranty_period = ?,
                    amc_period = ?,
                    payment_terms = ?,
                    manufacturer_types = ?,
                    supplying_capacity = ?,
                    mse_relaxation = ?,
                    similar_work_definition = ?,
                    annualized_value = ?,
                    escalation_clause = ?,
                    divisibility = ?,
                    performance_security = ?,
                    proposed_by = ?,
                    recommended_by = ?,
                    concurred_by = ?,
                    approved_by = ?,
                    amc_value = ?,
                    has_amc = ?,
                    correction_factor = ?,
                    o_m_value = ?,
                    o_m_period = ?,
                    has_om = ?,
                    additional_details = ?,
                    updated_at = CURRENT_TIMESTAMP
                WHERE id = ?
                ''', (
                    db_data['group_name'],
                    db_data['item_name'],
                    db_data['project_name'],
                    db_data['tender_description'],
                    db_data['pr_reference'],
                    db_data['tender_type'],
                    db_data['cec_estimate_incl_gst'],
                    db_data['cec_date'],
                    db_data['cec_estimate_excl_gst'],
                    db_data['budget_details'],
                    db_data['tender_platform'],
                    db_data['scope_of_work'],
                    db_data['contract_period_years'],
                    db_data['delivery_period'],
                    db_data['warranty_period'],
                    db_data['amc_period'],
                    db_data['payment_terms'],
                    db_data['manufacturer_types'],
                    db_data['supplying_capacity'],
                    db_data['mse_relaxation'],
                    db_data['similar_work_definition'],
                    db_data['annualized_value'],
                    db_data['escalation_clause'],
                    db_data['divisibility'],
                    db_data['performance_security'],
                    db_data['proposed_by'],
                    db_data['recommended_by'],
                    db_data['concurred_by'],
                    db_data['approved_by'],
                    db_data['amc_value'],
                    db_data['has_amc'],
                    db_data['correction_factor'],
                    db_data['o_m_value'],
                    db_data['o_m_period'],
                    db_data['has_om'],
                    db_data['additional_details'],
                    existing_record[0]
                ))
                message = "Data updated successfully!"
            else:
                # Insert new record
                cursor.execute('''
                INSERT INTO bqc_data (
                    user_id, ref_number, group_name, item_name, project_name, 
                    tender_description, pr_reference, tender_type, 
                    cec_estimate_incl_gst, cec_date, cec_estimate_excl_gst, 
                    budget_details, tender_platform, scope_of_work, 
                    contract_period_years, delivery_period, 
                    warranty_period, amc_period, payment_terms, 
                    manufacturer_types, supplying_capacity, mse_relaxation, 
                    similar_work_definition, annualized_value, escalation_clause, 
                    divisibility, performance_security, 
                    proposed_by, recommended_by, concurred_by, approved_by, 
                    amc_value, has_amc, correction_factor, o_m_value, o_m_period, has_om, additional_details
                ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                ''', (
                    self.user_id,
                    self.data['ref_number'],
                    db_data['group_name'],
                    db_data['item_name'],
                    db_data['project_name'],
                    db_data['tender_description'],
                    db_data['pr_reference'],
                    db_data['tender_type'],
                    db_data['cec_estimate_incl_gst'],
                    db_data['cec_date'],
                    db_data['cec_estimate_excl_gst'],
                    db_data['budget_details'],
                    db_data['tender_platform'],
                    db_data['scope_of_work'],
                    db_data['contract_period_years'],
                    db_data['delivery_period'],
                    db_data['warranty_period'],
                    db_data['amc_period'],
                    db_data['payment_terms'],
                    db_data['manufacturer_types'],
                    db_data['supplying_capacity'],
                    db_data['mse_relaxation'],
                    db_data['similar_work_definition'],
                    db_data['annualized_value'],
                    db_data['escalation_clause'],
                    db_data['divisibility'],
                    db_data['performance_security'],
                    db_data['proposed_by'],
                    db_data['recommended_by'],
                    db_data['concurred_by'],
                    db_data['approved_by'],
                    db_data['amc_value'],
                    db_data['has_amc'],
                    db_data['correction_factor'],
                    db_data['o_m_value'],
                    db_data['o_m_period'],
                    db_data['has_om'],
                    db_data['additional_details']
                ))
                message = "Data saved successfully!"
            
            conn.commit()
            conn.close()
            
            self.statusBar().showMessage(message)
            QMessageBox.information(self, "Success", message)
            
        except Exception as e:
            logger.error(f"Error saving data: {str(e)}")
            QMessageBox.critical(self, "Error", f"Failed to save data: {str(e)}")
        finally:
            progress.close()
    
    def load_data(self):
        """Load saved data from database"""
        dialog = LoadDataDialog(self.user_id, self)
        if dialog.exec_() == QDialog.Accepted and dialog.selected_id:
            # Show progress dialog
            progress = QProgressDialog("Loading data...", "Cancel", 0, 0, self)
            progress.setWindowModality(Qt.WindowModal)
            progress.show()
            
            try:
                conn = sqlite3.connect(DB_PATH)
                cursor = conn.cursor()
                
                cursor.execute('''
                SELECT * FROM bqc_data WHERE id = ?
                ''', (dialog.selected_id,))
                
                record = cursor.fetchone()
                
                if record:
                    # Map database columns to data keys
                    columns = [desc[0] for desc in cursor.description]
                    record_dict = dict(zip(columns, record))
                    
                    # Update form data
                    self.data['ref_number'] = record_dict['ref_number'] or ''
                    self.data['group_name'] = record_dict['group_name'] or '1 - Materials'
                    self.data['item_name'] = record_dict['item_name'] or ''
                    self.data['project_name'] = record_dict['project_name'] or ''
                    self.data['tender_description'] = record_dict['tender_description'] or ''
                    self.data['pr_reference'] = record_dict['pr_reference'] or ''
                    self.data['tender_type'] = record_dict['tender_type'] or 'Goods'
                    self.data['cec_estimate_incl_gst'] = record_dict['cec_estimate_incl_gst'] or 0
                    
                    # Handle date
                    if record_dict['cec_date']:
                        try:
                            self.data['cec_date'] = datetime.datetime.strptime(record_dict['cec_date'], "%Y-%m-%d").date()
                        except:
                            self.data['cec_date'] = datetime.date.today()
                    else:
                        self.data['cec_date'] = datetime.date.today()
                    
                    self.data['cec_estimate_excl_gst'] = record_dict['cec_estimate_excl_gst'] or 0
                    self.data['budget_details'] = record_dict['budget_details'] or ''
                    self.data['tender_platform'] = record_dict['tender_platform'] or 'GeM'
                    self.data['scope_of_work'] = record_dict['scope_of_work'] or ''
                    
                    # Handle contract period
                    if 'contract_period_years' in record_dict and record_dict['contract_period_years'] is not None:
                        self.data['contract_period_years'] = record_dict['contract_period_years']
                    elif 'contract_period_months_old' in record_dict and record_dict['contract_period_months_old'] is not None:
                        # Convert months to years
                        self.data['contract_period_years'] = record_dict['contract_period_months_old'] / 12.0
                    elif 'contract_period_months' in record_dict and record_dict['contract_period_months'] is not None:
                        # Convert months to years
                        self.data['contract_period_years'] = record_dict['contract_period_months'] / 12.0
                    else:
                        self.data['contract_period_years'] = 1.0  # Default value
                    
                    self.data['delivery_period'] = record_dict['delivery_period'] or ''
                    self.data['warranty_period'] = record_dict['warranty_period'] or ''
                    self.data['amc_period'] = record_dict['amc_period'] or ''
                    self.data['payment_terms'] = record_dict['payment_terms'] or ''
                    
                    # Handle manufacturer types
                    try:
                        if record_dict['manufacturer_types']:
                            self.data['manufacturer_types'] = json.loads(record_dict['manufacturer_types'])
                        else:
                            self.data['manufacturer_types'] = ['Original Equipment Manufacturer']
                    except:
                        self.data['manufacturer_types'] = ['Original Equipment Manufacturer']
                    
                    self.data['supplying_capacity'] = record_dict['supplying_capacity'] or 30
                    self.data['mse_relaxation'] = bool(record_dict['mse_relaxation'])
                    self.data['similar_work_definition'] = record_dict['similar_work_definition'] or ''
                    self.data['annualized_value'] = record_dict['annualized_value'] or 0
                    self.data['escalation_clause'] = record_dict['escalation_clause'] or ''
                    self.data['divisibility'] = record_dict['divisibility'] or 'Non-Divisible'
                    self.data['performance_security'] = record_dict['performance_security'] or 5
                    self.data['proposed_by'] = record_dict['proposed_by'] or 'XXXXX'
                    self.data['recommended_by'] = record_dict['recommended_by'] or 'XXXXX'
                    self.data['concurred_by'] = record_dict['concurred_by'] or 'Rajesh J.'
                    self.data['approved_by'] = record_dict['approved_by'] or 'Kani Amudhan N.'
                    self.data['amc_value'] = record_dict['amc_value'] or 0
                    self.data['has_amc'] = bool(record_dict['has_amc'])
                    self.data['correction_factor'] = record_dict['correction_factor'] or 0
                    self.data['o_m_value'] = record_dict.get('o_m_value', 0) or 0
                    self.data['o_m_period'] = record_dict.get('o_m_period', '') or ''
                    self.data['has_om'] = bool(record_dict.get('has_om', 0))
                    self.data['additional_details'] = record_dict.get('additional_details', '') or ''
                    
                    # Update form fields
                    self.ref_number_input.setText(self.data['ref_number'])
                    
                    # Find and set group combo
                    for i in range(self.group_combo.count()):
                        if self.group_combo.itemText(i).startswith(self.data['group_name'].split(' - ')[0]):
                            self.group_combo.setCurrentIndex(i)
                            break
                    
                    self.tender_description_input.setText(self.data['tender_description'])
                    self.pr_reference_input.setText(self.data['pr_reference'])
                    self.tender_type_combo.setCurrentText(self.data['tender_type'])
                    self.cec_incl_gst_input.setValue(self.data['cec_estimate_incl_gst'])
                    self.cec_date_input.setDate(QDate.fromString(self.data['cec_date'].isoformat(), "yyyy-MM-dd"))
                    self.cec_excl_gst_input.setValue(self.data['cec_estimate_excl_gst'])
                    self.budget_details_input.setText(self.data['budget_details'])
                    self.platform_combo.setCurrentText(self.data['tender_platform'])
                    self.scope_of_work_input.setText(self.data['scope_of_work'])
                    self.contract_period_years_input.setValue(self.data['contract_period_years'])
                    self.amc_period_input.setText(self.data['amc_period'])
                    self.payment_terms_input.setText(self.data['payment_terms'])
                    self.delivery_period_input.setText(self.data['delivery_period'])
                    self.warranty_period_input.setText(self.data['warranty_period'])
                    self.escalation_clause_input.setText(self.data['escalation_clause'])
                    self.additional_details_input.setText(self.data['additional_details'])
                    self.divisibility_combo.setCurrentText(self.data['divisibility'])
                    self.performance_security_input.setValue(self.data['performance_security'])
                    self.proposed_by_input.setText(self.data['proposed_by'])
                    self.recommended_by_input.setText(self.data['recommended_by'])
                    self.concurred_by_input.setText(self.data['concurred_by'])
                    self.approved_by_input.setText(self.data['approved_by'])
                    self.amc_value_input.setValue(self.data['amc_value'])
                    self.has_amc_checkbox.setChecked(self.data['has_amc'])
                    self.correction_factor_input.setValue(self.data['correction_factor'])
                    
                    # Update manufacturer checkboxes
                    for i, cb in enumerate(self.manufacturer_checkboxes):
                        cb.setChecked(MANUFACTURER_TYPES[i] in self.data['manufacturer_types'])
                    
                    # Update MSE relaxation
                    self.mse_relaxation_checkbox.setChecked(self.data['mse_relaxation'])
                    
                    # Update supplying capacity
                    self.supplying_capacity_input.setValue(self.data['supplying_capacity'])
                    self.update_supplying_capacity(self.data['supplying_capacity'])
                    
                    # Update calculated values
                    self.update_calculated_values()
                    
                    # Update visibility
                    self.update_amc_visibility()
                    self.update_bqc_tab_visibility()
                    self.update_divisibility(self.data['divisibility'])
                    
                    self.statusBar().showMessage("Data loaded successfully!")
                    QMessageBox.information(self, "Success", "Data loaded successfully!")
                
                conn.close()
                
            except Exception as e:
                logger.error(f"Error loading data: {str(e)}")
                QMessageBox.critical(self, "Error", f"Failed to load data: {str(e)}")
            finally:
                progress.close()
    
    def clear_form(self):
        # Reset data to defaults
        self.data = {
            'ref_number': '',
            'group_name': '1 - Materials',
            'item_name': '',
            'project_name': '',
            'tender_description': '',
            'pr_reference': '',
            'tender_type': 'Goods',
            'cec_estimate_incl_gst': 0,
            'cec_date': datetime.date.today(),
            'cec_estimate_excl_gst': 0,
            'budget_details': '',
            'tender_platform': 'GeM',
            'scope_of_work': '',
            'contract_period_years': 1,
            'delivery_period': '',
            'warranty_period': '',
            'amc_period': '',
            'payment_terms': '',
            'manufacturer_types': ['Original Equipment Manufacturer'],
            'supplying_capacity': 30,
            'mse_relaxation': False,
            'similar_work_definition': '',
            'annualized_value': 0,
            'escalation_clause': '',
            'divisibility': 'Non-Divisible',
            'performance_security': 5,
            'proposed_by': 'XXXXX',
            'recommended_by': 'XXXXX',
            'concurred_by': 'Rajesh J.',
            'approved_by': 'Kani Amudhan N.',
            'amc_value': 0,
            'has_amc': False,
            'correction_factor': 0,
            'o_m_value': 0,
            'o_m_period': '',
            'has_om': False,
            'additional_details': ''
        }
        
        # Reset all form fields
        self.ref_number_input.setText(self.data['ref_number'])
        self.group_combo.setCurrentIndex(0)
        self.tender_description_input.setText(self.data['tender_description'])
        self.pr_reference_input.setText(self.data['pr_reference'])
        self.tender_type_combo.setCurrentText(self.data['tender_type'])
        self.cec_incl_gst_input.setValue(self.data['cec_estimate_incl_gst'])
        self.cec_date_input.setDate(QDate.currentDate())
        self.cec_excl_gst_input.setValue(self.data['cec_estimate_excl_gst'])
        self.budget_details_input.setText(self.data['budget_details'])
        self.platform_combo.setCurrentText(self.data['tender_platform'])
        self.scope_of_work_input.setText(self.data['scope_of_work'])
        self.contract_period_years_input.setValue(self.data['contract_period_years'])
        self.amc_period_input.setText(self.data['amc_period'])
        self.payment_terms_input.setText(self.data['payment_terms'])
        self.delivery_period_input.setText(self.data['delivery_period'])
        self.warranty_period_input.setText(self.data['warranty_period'])
        self.escalation_clause_input.setText(self.data['escalation_clause'])
        self.additional_details_input.setText(self.data['additional_details'])
        self.divisibility_combo.setCurrentText(self.data['divisibility'])
        self.performance_security_input.setValue(self.data['performance_security'])
        self.proposed_by_input.setText(self.data['proposed_by'])
        self.recommended_by_input.setText(self.data['recommended_by'])
        self.concurred_by_input.setText(self.data['concurred_by'])
        self.approved_by_input.setText(self.data['approved_by'])
        self.amc_value_input.setValue(self.data['amc_value'])
        self.has_amc_checkbox.setChecked(self.data['has_amc'])
        self.correction_factor_input.setValue(self.data['correction_factor'])
        
        # Reset manufacturer checkboxes
        for i, cb in enumerate(self.manufacturer_checkboxes):
            cb.setChecked(MANUFACTURER_TYPES[i] in self.data['manufacturer_types'])
        
        # Reset MSE relaxation
        self.mse_relaxation_checkbox.setChecked(self.data['mse_relaxation'])
        self.mse_calculation_label.setVisible(False)
        
        # Reset supplying capacity
        self.supplying_capacity_input.setValue(self.data['supplying_capacity'])
        self.update_supplying_capacity(self.data['supplying_capacity'])
        
        # Update calculated values
        self.update_calculated_values()
        
        # Update visibility
        self.update_amc_visibility()
        self.update_bqc_tab_visibility()
        self.update_divisibility(self.data['divisibility'])
        
        self.statusBar().showMessage("Form cleared successfully!")
    
    def generate_document(self):
        # Validate form data
        is_valid, errors = validate_input(self.data)
        
        if not is_valid:
            error_msg = "Please fix the following errors:\n\n" + "\n".join([f"• {error}" for error in errors])
            QMessageBox.warning(self, "Validation Error", error_msg)
            return
        
        # Show progress dialog
        progress = QProgressDialog("Generating document...", "Cancel", 0, 0, self)
        progress.setWindowModality(Qt.WindowModal)
        progress.show()
        
        # Generate the document
        doc_bytes = generate_bqc_document(self.data)
        
        progress.close()
        
        if doc_bytes:
            # Ask user where to save the file
            file_name = f"BQC_{self.data['ref_number']}_{datetime.date.today().strftime('%Y%m%d')}.docx"
            file_path, _ = QFileDialog.getSaveFileName(
                self, "Save Document", file_name, "Word Documents (*.docx)"
            )
            
            if file_path:
                try:
                    with open(file_path, 'wb') as f:
                        f.write(doc_bytes.getvalue())
                    QMessageBox.information(self, "Success", "Document generated and saved successfully!")
                    self.statusBar().showMessage(f"Document saved to {file_path}")
                except Exception as e:
                    QMessageBox.critical(self, "Error", f"Failed to save document: {str(e)}")
                    logger.error(f"Error saving document: {str(e)}", exc_info=True)
        else:
            QMessageBox.critical(self, "Error", "Failed to generate document. Please check the logs for details.")

# Helper functions
def calculate_emd(estimated_value: float, tender_type: str) -> float:
    """Calculate EMD amount based on estimated value and tender type"""
    if estimated_value < 50:
        return 0
    
    for threshold, emd in EMD_THRESHOLDS:
        if estimated_value <= threshold:
            # Special case for Goods/Services between 50-100
            if threshold == 100 and tender_type in ['Goods', 'Services']:
                return 0
            return emd
    return 20

def validate_input(data: Dict) -> Tuple[bool, List[str]]:
    """Validate user inputs and return (is_valid, error_messages)"""
    errors = []
    
    # Required fields validation
    required_fields = [
        'ref_number', 'tender_description',
        'pr_reference', 'budget_details', 'scope_of_work'
    ]
    
    for field in required_fields:
        if not data.get(field):
            errors.append(f"Field '{field.replace('_', ' ').title()}' is required")
    
    # Numeric validation
    if data['cec_estimate_incl_gst'] <= 0:
        errors.append("CEC Estimate (incl. GST) must be greater than 0")
    
    if data['cec_estimate_excl_gst'] <= 0:
        errors.append("CEC Estimate (excl. GST) must be greater than 0")
    
    if data['cec_estimate_incl_gst'] < data['cec_estimate_excl_gst']:
        errors.append("CEC Estimate (incl. GST) must be greater than or equal to CEC Estimate (excl. GST)")
    
    # Contract period validation
    if data['contract_period_years'] <= 0:
        errors.append("Contract Period (Years) must be greater than 0")
    
    # Goods-specific validation
    if data['tender_type'] == 'Goods':
        if not data.get('delivery_period'):
            errors.append("Delivery Period is required for Goods tenders")
        if not data.get('warranty_period'):
            errors.append("Warranty Period is required for Goods tenders")
    
    # Service/Works-specific validation
    if data['tender_type'] in ['Service', 'Works']:
        if not data.get('similar_work_definition'):
            errors.append("Definition of Similar Work is required for Service/Works tenders")
    
    return len(errors) == 0, errors

def generate_bqc_document(data: Dict) -> Optional[BytesIO]:
    """Generate the BQC document in .docx format based on user inputs"""
    try:
        doc = Document()
        
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)
        
        # Calculate EMD amount
        emd_amount = calculate_emd(data['cec_estimate_excl_gst'], data['tender_type'])
        
        # Calculate turnover requirement with refined logic
        base_percentage = 0.3
        
        # Apply correction factor when divisible for all tender types
        if data['divisibility'] == 'Divisible':
            base_percentage = 0.3 * (1 + data['correction_factor'])
        
        # Calculate turnover requirement considering AMC value
        maintenance_value = 0
        
        if data['has_amc']:
            maintenance_value = data['amc_value']
        
        if maintenance_value > 0:
            # Calculate as base_percentage of (CEC_incl_gst - maintenance_value)
            turnover_requirement = base_percentage * (data['cec_estimate_incl_gst'] - maintenance_value)
        else:
            # Calculate as base_percentage of CEC
            turnover_requirement = base_percentage * data['cec_estimate_excl_gst']
        
        # Convert to Crore for display
        turnover_requirement_crore = turnover_requirement / 100
        
        # Calculate experience requirements for Service/Works
        if data['tender_type'] in ['Service', 'Works']:
            # Apply correction factor if divisible
            if data['divisibility'] == 'Divisible':
                correction_factor = data['correction_factor']
                option_a_percent = 0.4 * (1 + correction_factor)
                option_b_percent = 0.5 * (1 + correction_factor)
                option_c_percent = 0.8 * (1 + correction_factor)
            else:
                option_a_percent = 0.4
                option_b_percent = 0.5
                option_c_percent = 0.8
            
            option_a_value = option_a_percent * data['cec_estimate_incl_gst']
            option_b_value = option_b_percent * data['cec_estimate_incl_gst']
            option_c_value = option_c_percent * data['cec_estimate_incl_gst']
        
        # Calculate Annualized Estimated Value
        if data['contract_period_years'] > 0:
            annualized_value = data['cec_estimate_excl_gst'] / data['contract_period_years']
        else:
            annualized_value = data['annualized_value']
        
        # Calculate supplying capacity (30% of base value)
        base_supplying_capacity = data['supplying_capacity']
        calculated_supplying_capacity = int(base_supplying_capacity * 0.3)
        
        # Apply MSE relaxation if applicable
        if data['mse_relaxation']:
            final_supplying_capacity = int(calculated_supplying_capacity * 0.85)  # 15% relaxation
        else:
            final_supplying_capacity = calculated_supplying_capacity
        
        # Header section with table for proper alignment
        table = doc.add_table(rows=3, cols=2)
        table.style = 'Table Grid'
        
        # Row 0: Ref and Date
        table.cell(0, 0).text = f"Ref: {data['ref_number']}"
        table.cell(0, 1).text = f"Date: {datetime.date.today().strftime('%d/%m/%Y')}"
        
        # Row 1: Note To
        note_to_cell = table.cell(1, 0).merge(table.cell(1, 1))
        note_to_cell.text = "NOTE TO: CHIEF PROCUREMENT OFFICER, CPO (M)/ PROCUREMENT LEADER GROUP XX"
        
        # Row 2: Subject
        subject_cell = table.cell(2, 0).merge(table.cell(2, 1))
        
        # Determine subject text based on tender type
        if data['tender_type'] == 'Goods':
            subject_text = f"SUPPLY OF ITEMS FOR '{data['tender_description']}'"
        else:
            subject_text = f"JOB OF CONSTRUCTION OF '{data['tender_description']}'"
        
        subject_cell.text = f"SUBJECT: {subject_text}: APPROVAL OF BID QUALIFICATION CRITERIA AND FLOATING OF OPEN DOMESTIC TENDER."
        
        # Add some space after the table
        doc.add_paragraph()
        
        # Section 1: PREAMBLE
        p = doc.add_paragraph()
        p.add_run("1.\tPREAMBLE").bold = True
        
        # Create table for PREAMBLE
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Table Grid'
        
        # Fill table with data
        table.cell(0, 0).text = "Tender Description"
        table.cell(0, 1).text = data['tender_description']
        
        table.cell(1, 0).text = "PR reference/ Email reference"
        table.cell(1, 1).text = data['pr_reference']
        
        table.cell(2, 0).text = "Type of Tender"
        table.cell(2, 1).text = data['tender_type']
        
        table.cell(3, 0).text = "CEC estimate (incl. of GST)/ Date"
        table.cell(3, 1).text = f"{data['cec_estimate_incl_gst']} / {data['cec_date'].strftime('%d/%m/%Y')}"
        
        table.cell(4, 0).text = "CEC estimate exclusive of GST"
        table.cell(4, 1).text = str(data['cec_estimate_excl_gst'])
        
        p = doc.add_paragraph()
        p.add_run("Tender Platform – ").bold = True
        p.add_run(data['tender_platform'])
        
        # Add some space
        doc.add_paragraph("\t")
        
        # Section 2: BRIEF SCOPE OF WORK/ SUPPLY ITEMS
        p = doc.add_paragraph()
        p.add_run("2.\tBRIEF SCOPE OF WORK/ SUPPLY ITEMS").bold = True
        
        # Create table for BRIEF SCOPE with dynamic rows
        # Base rows: scope, contract period, payment terms
        base_rows = 3
        
        # Additional rows for Goods
        if data['tender_type'] == 'Goods':
            additional_rows = 2  # delivery and warranty
        else:
            additional_rows = 0
        
        # Additional rows for maintenance (AMC)
        if data['has_amc']:
            additional_rows += 2  # period and value
        
        total_rows = base_rows + additional_rows
        table = doc.add_table(rows=total_rows, cols=2)
        table.style = 'Table Grid'
        
        # Fill table with data
        row_idx = 0
        
        # Row 0: Brief Scope of Work / Supply Items
        table.cell(row_idx, 0).text = "Brief Scope of Work / Supply Items"
        table.cell(row_idx, 1).text = data['scope_of_work']
        row_idx += 1
        
        # Row 1: Contract Period
        table.cell(row_idx, 0).text = "Contract Period"
        table.cell(row_idx, 1).text = f"{data['contract_period_years']} years"
        row_idx += 1
        
        # Additional rows for Goods
        if data['tender_type'] == 'Goods':
            # Row: Delivery Period
            table.cell(row_idx, 0).text = "Delivery Period of the Item"
            table.cell(row_idx, 1).text = data['delivery_period']
            row_idx += 1
            
            # Row: Warranty Period
            table.cell(row_idx, 0).text = "Warranty Period"
            table.cell(row_idx, 1).text = data['warranty_period']
            row_idx += 1
        
        # Additional rows for maintenance (AMC)
        if data['has_amc']:
            # Row: AMC Period
            table.cell(row_idx, 0).text = "AMC/ CAMC Period (No. of Years)"
            table.cell(row_idx, 1).text = data['amc_period']
            row_idx += 1
            
            # Row: AMC Value
            table.cell(row_idx, 0).text = "AMC/ CAMC Value (Lakhs)"
            table.cell(row_idx, 1).text = str(data['amc_value'])
            row_idx += 1
        
        # Last row: Payment Terms
        table.cell(row_idx, 0).text = "Payment Terms (if different from standard terms i.e within 30 days)"
        table.cell(row_idx, 1).text = data['payment_terms']
        
        # Add some space
        doc.add_paragraph("\t")
        
        # Section 3: BID QUALIFICATION CRITERIA (BQC)
        p = doc.add_paragraph()
        p.add_run("3.\tBID QUALIFICATION CRITERIA (BQC)").bold = True
        p = doc.add_paragraph("BPCL would like to qualify vendors for undertaking the above work as indicated in the brief scope. Detailed bid qualification criteria for short listing vendors shall be as follows:")
        
        # Technical Criteria based on tender type
        if data['tender_type'] == 'Goods':
            p = doc.add_paragraph()
            p.add_run("3.1\tTECHNICAL CRITERIA").bold = True
            p = doc.add_paragraph()
            p.add_run("3.1.1. For GOODS:").bold = True
            
            # Manufacturing Capability
            p = doc.add_paragraph()
            p.add_run("a) Manufacturing Capability:").bold = True
            p = doc.add_paragraph("Bidder* should be ")
            
            # Add selected manufacturer types
            manufacturer_types = ", ".join(data['manufacturer_types'])
            p.add_run(manufacturer_types).italic = True
            p.add_run(" of the item being tendered. ")
            
            p = doc.add_paragraph("*The definition of bidder is the entity which has a unique PAN (Permanent Account Number). All documents should be in the name of the bidder only (except in cases where the bidder is allowed to take the technical credentials of their OEM). Documents in the name of any legal entity other than the bidder, as defined above, shall not be accepted")
            
            p = doc.add_paragraph("(Explanatory Note: This is to ensure that the items are procured from the established Manufacturers, thereby ensuring right quality and price. OEMs are generally required to quote/participate in the tender. However, on case to case basis traders, authorized dealers, distributors, etc. may be considered depending on market scenario and volume of the tender.)")
            
            # Supplying Capacity
            p = doc.add_paragraph()
            p.add_run("b) Supplying Capacity:").bold = True
            
            p = doc.add_paragraph(f"The bidder shall have experience of having successfully supplied minimum of {final_supplying_capacity} quantity of the annualized estimated quantity in any 12 continuous months during last 7 years in India or abroad, ending on last day of the month previous to the one in which tender is invited.")
            
            if data['mse_relaxation']:
                p = doc.add_paragraph("For MSE bidders Relaxation of 15% on the supplying capacity shall be given as per Corp. Finance Circular MA.TEC.POL.CON.3A dated 26.10.2020.")
            
            p = doc.add_paragraph("(Explanatory Note: The above criterion will demonstrate vendor's capability to manufacture the tendered item satisfactorily Based on market search 30 quantity of the annualized estimated quantity may be relaxed for widening the competition for specific items)")
            
            p = doc.add_paragraph("*The definition of bidder is the entity which has a unique PAN (Permanent Account Number). All documents should be in the name of the bidder only (except in cases where the bidder is allowed to take the technical credentials of their OEM). Documents in the name of any legal entity other than the bidder, as defined above, shall not be accepted.")
            
        else:  # Service or Works
            p = doc.add_paragraph()
            p.add_run("3.1.2. BQC/PQC for Procurement of Works and Services:").bold = True
            
            p = doc.add_paragraph()
            p.add_run("I) Experience / Past performance / Technical Capability:").bold = True
            
            p = doc.add_paragraph("The bidder# should have experience of having successfully completed similar works during last 7 years ending last day of month previous to the one in which tender is floated should be either of the following: -")
            
            p = doc.add_paragraph(f"a. Three similar completed works each costing not less than the amount equal to {option_a_percent*100:.0f}% of the estimated cost.")
            p = doc.add_paragraph("or")
            p = doc.add_paragraph(f"b. Two similar completed works each costing not less than the amount equal to {option_b_percent*100:.0f}% of the estimated cost.")
            p = doc.add_paragraph("or")
            p = doc.add_paragraph(f"c. One similar completed work costing not less than the amount equal to {option_c_percent*100:.0f}% of the estimated cost.")
            
            p = doc.add_paragraph(f'Definition of "similar work" should be clearly defined: {data["similar_work_definition"]}')
            
            p = doc.add_paragraph("# In case of Service contracts the term bidder may be suitably modified to take care of OEMs/ System Integrators/ Authorised Channel Partner etc.")
            p = doc.add_paragraph("*The definition of bidder is the entity which has a unique PAN (Permanent Account Number). All documents should be in the name of the bidder only (except in cases where the bidder is allowed to take the technical credentials of their OEM). Documents in the name of any legal entity other than the bidder, as defined above, shall not be accepted.")
            
            p = doc.add_paragraph("[ The % mentioned for meeting the similar works value can be amended considering the market inputs for specific items]")
            p = doc.add_paragraph("[In addition to above, additional criteria may be incorporated according to the requirement of the Project].")
        
        # Financial Criteria
        p = doc.add_paragraph()
        p.add_run("3.2\tFINANCIAL CRITERIA").bold = True
        
        p = doc.add_paragraph()
        p.add_run("3.2.1\tAVERAGE ANNUAL TURNOVER").bold = True
        
        # Only show the turnover requirement without AMC calculation details
        p.add_run(f"The average annual turnover of the Bidder for last three audited accounting years shall be equal to or more than {base_percentage*100:.0f}% of the annualized estimated value in Rs. {turnover_requirement_crore:.2f} Crore.")
        
        p = doc.add_paragraph("Explanatory Notes:")
        p = doc.add_paragraph("i. Average annual turnover values in-line with CTE Office Memorandum No. 12-02-1-CTE-6 dated 17th Dec 2002.")
        
        p = doc.add_paragraph()
        p.add_run("3.2.2\tNET WORTH").bold = True
        p.add_run("The bidder should have positive net worth as per the latest audited financial statement.")
        
        p = doc.add_paragraph("Documents Required: Please refer the ITB (Instruction to Bidders) which mentions the documents to be submitted by bidders for meeting the above Technical and Financial criteria.")
        
        # BIDS MAY BE SUBMITTED BY
        p = doc.add_paragraph()
        p.add_run("3.3\tBIDS MAY BE SUBMITTED BY").bold = True
        
        p = doc.add_paragraph()
        p.add_run("3.3.1\t").bold = True
        p.add_run("An entity (domestic bidder) should have completed 3 financial years of existence as on original due date of tender since date of commencement of business and shall fulfil each BQC eligibility criteria as mentioned above.")
        
        p = doc.add_paragraph()
        p.add_run("3.3.2\t").bold = True
        p.add_run("JV/Consortium bids will not be accepted (i.e. Qualification on the strength of the JV Partners/Consortium Members /Subsidiaries / Group members will not be accepted)")
        
        # ESCALATION/ DE-ESCALATION CLAUSE
        if data['escalation_clause']:
            p = doc.add_paragraph()
            p.add_run("4. ESCALATION/ DE-ESCALATION CLAUSE: ").bold = True
            p.add_run(data['escalation_clause'])
        
        # ADDITIONAL DETAILS
        if data['additional_details']:
            p = doc.add_paragraph()
            p.add_run("5. ADDITIONAL DETAILS: ").bold = True
            p.add_run(data['additional_details'])
        
        # EARNEST MONEY DEPOSIT (EMD)
        p = doc.add_paragraph()
        p.add_run("6.\tEARNEST MONEY DEPOSIT (EMD)").bold = True
        
        # Format EMD amount (0 for Nil)
        if emd_amount == 0:
            emd_text = "Nil"
        else:
            emd_text = f"{emd_amount} Lacs"
        
        p = doc.add_paragraph(f"Bidders are required to provide Earnest Money Deposit equivalent to Rs. {emd_text} for the tender.")
        
        p = doc.add_paragraph("EMD exemption shall be as per General Terms & Conditions of GeM (applicable for GeM tenders)/ MSE policy")
        
        p = doc.add_paragraph("Explanatory Note:")
        doc.add_paragraph("Procurement Group to justify the EMD amount as per Guidelines.")
        
        # Performance Security
        p = doc.add_paragraph()
        p.add_run("7. Performance Security (if at variance with the ITB clause):").bold = True
        
        # Determine standard performance security percentage
        if data['tender_type'] in ['Goods', 'Services']:
            standard_ps = 5
        else:  # Works
            standard_ps = 10
        
        if data['performance_security'] != standard_ps:
            p.add_run(f"Performance Security {data['performance_security']}% (approved by the competent authority).")
            
            # Add note about non-standard performance security
            p = doc.add_paragraph()
            p.add_run("Note: ").bold = True
            p.add_run(f"The performance security percentage of {data['performance_security']}% is different from the standard percentage of {standard_ps}% for {data['tender_type']} tenders. This has been approved by the competent authority.")
        else:
            p.add_run("Performance Security as per standard terms (5% for Goods & Services, 10% for Works).")
        
        # Other Points
        p = doc.add_paragraph("Other Points which may be taken into consideration while framing BQC:")
        p = doc.add_paragraph("1) Any guidelines from company, govt., industry tender practices (in case of industry tenders) etc. shall need to be followed superseding the above criteria as applicable.")
        p = doc.add_paragraph("2) Any services rendered by the vendor after due supply of the goods like AMC/CAMC after warranty period, servicing, etc. needs to be appropriately excluded while fixing the qualification (technical) criteria.")
        
        if data['has_amc']:
            doc.add_paragraph("3) Where, the tender involves Annual Maintenance Contract (AMC) or Comprehensive Annual Maintenance Contract (CAMC), the estimated cost towards AMC/CAMC shall be excluded while arriving at the financial criteria (Annual Turnover) for the tender.")
        else:
            doc.add_paragraph("3) Where, the tender involves Annual Maintenance Contract (AMC) or Comprehensive Annual Maintenance Contract (CAMC), the estimated cost towards AMC/CAMC shall be excluded while arriving at the financial criteria (Annual Turnover) for the tender.")
        
        p = doc.add_paragraph("4) Additional qualification criteria may be built upon depending on the situation on case to case basis.")
        p = doc.add_paragraph("5) During first time procurement of any goods/services by CPO (M), inputs from User SBU/Entity should be taken. However same may be taken for subsequent procurements also to the extent possible and depending on the complexity of the procurement.")
        
        # APPROVAL REQUIRED
        p = doc.add_paragraph()
        p.add_run("8. APPROVAL REQUIRED").bold = True
        
        p = doc.add_paragraph("In view of above, approval is requested for the Supply of ")
        p.add_run(f"items/ job -{data['tender_description']}").bold = True
        p.add_run(" for:")
        
        p = doc.add_paragraph()
        p.add_run("i.\t").bold = True
        p.add_run("Bid Qualification Criteria as per Sr. No. 3, as per Clause 13.8 of Guidelines for procurement of Goods and Contract Services.")
        
        p = doc.add_paragraph()
        p.add_run("ii.\t").bold = True
        p.add_run("Inviting bids (two-part bid) through a Domestic Open Tender.")
        
        p = doc.add_paragraph()
        p.add_run("iii.\t").bold = True
        p.add_run("Earnest Money Deposit as per Sr. No. 6 above./ Performance Security as per Sr. No. 7 (if applicable)")
        
        # Add some space
        doc.add_paragraph("\t")
        
        # Approval section
        p = doc.add_paragraph()
        p.add_run("Proposed by").bold = True
        p = doc.add_paragraph(f"{data['proposed_by']}, Procurement Manager (CPO Mktg.)")
        
        p = doc.add_paragraph()
        p.add_run("Recommended by").bold = True
        p = doc.add_paragraph(f"{data['recommended_by']}, Procurement Leader (CPO Mktg.)")
        
        p = doc.add_paragraph()
        p.add_run("Concurred by").bold = True
        p = doc.add_paragraph(f"{data['concurred_by']}, GM Finance (CPO Mktg.)")
        
        p = doc.add_paragraph()
        p.add_run("Approved by").bold = True
        p = doc.add_paragraph(f"{data['approved_by']}, Chief Procurement Officer, (CPO Mktg.)")
        
        # Save the document to a bytes buffer
        doc_bytes = BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        
        logger.info(f"Successfully generated BQC document for ref: {data['ref_number']}")
        return doc_bytes
    
    except Exception as e:
        logger.error(f"Error generating BQC document: {str(e)}", exc_info=True)
        return None

if __name__ == "__main__":
    # Always check and setup database at startup
    if not setup_database():
        QMessageBox.critical(None, "Error", "Failed to initialize database. Application will exit.")
        sys.exit(1)
    
    app = QApplication(sys.argv)
    
    # Show login dialog
    login_dialog = LoginDialog()
    if login_dialog.exec_() == QDialog.Accepted:
        window = BQCGeneratorApp(login_dialog.user_id)
        window.show()
        sys.exit(app.exec_())
    else:
        sys.exit(0)
