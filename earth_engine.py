import streamlit as st
import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from io import BytesIO
import logging
import os
from typing import Dict, List, Tuple, Optional

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler("bqc_generator.log"),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# Constants
EMD_THRESHOLDS = [
    (50, 0),
    (100, 1),
    (500, 2.5),
    (1000, 5),
    (1500, 7.5),
    (2500, 10),
    (float('inf'), 20)
]

GROUP_OPTIONS = {
    '1': 'Materials',
    '2': 'Machinery & Equipment',
    '3': 'Projects',
    '4': 'Engineering Services',
    '5': 'Information Technology',
    '6': 'Marketing',
    '7': 'Human Resources',
    '8': 'Finance & Accounts',
    '9': 'Legal'
}

TENDER_TYPES = ["Goods", "Service", "Works"]
MANUFACTURER_TYPES = [
    "Original Equipment Manufacturer", 
    "Authorized Channel Partner", 
    "Authorized Agent", 
    "Dealer", 
    "Authorized Distributor"
]

EVALUATION_METHODS = [
    "Overall Lowest", 
    "Schedule wise Lowest", 
    "% bidding", 
    "LOT wise lowest", 
    "Least Cash Outflow", 
    "QCBS", 
    "Reverse auction"
]

DIVISIBILITY_OPTIONS = ["Non-Divisible", "Divisible"]
PLATFORM_OPTIONS = ["GeM", "E-procurement"]

def calculate_emd(estimated_value: float, tender_type: str) -> float:
    """Calculate EMD amount based on estimated value and tender type"""
    if estimated_value < 50:
        return 0
    
    for threshold, emd in EMD_THRESHOLDS:
        if estimated_value <= threshold:
            # Special case for Goods/Services between 50-100
            if threshold == 100 and tender_type in ['Goods', 'Services']:
                return 0
            return emd
    return 20

def validate_input(data: Dict) -> Tuple[bool, List[str]]:
    """Validate user inputs and return (is_valid, error_messages)"""
    errors = []
    
    # Required fields validation
    required_fields = [
        'ref_number', 'item_name', 'project_name', 'tender_description',
        'pr_reference', 'budget_details', 'scope_of_work', 'contract_period'
    ]
    
    for field in required_fields:
        if not data.get(field):
            errors.append(f"Field '{field.replace('_', ' ').title()}' is required")
    
    # Numeric validation
    if data['cec_estimate_incl_gst'] <= 0:
        errors.append("CEC Estimate (incl. GST) must be greater than 0")
    
    if data['cec_estimate_excl_gst'] <= 0:
        errors.append("CEC Estimate (excl. GST) must be greater than 0")
    
    if data['cec_estimate_incl_gst'] < data['cec_estimate_excl_gst']:
        errors.append("CEC Estimate (incl. GST) must be greater than or equal to CEC Estimate (excl. GST)")
    
    # Goods-specific validation
    if data['tender_type'] == 'Goods':
        if not data.get('delivery_period'):
            errors.append("Delivery Period is required for Goods tenders")
        if not data.get('warranty_period'):
            errors.append("Warranty Period is required for Goods tenders")
    
    # Service/Works-specific validation
    if data['tender_type'] in ['Service', 'Works']:
        if not data.get('similar_work_definition'):
            errors.append("Definition of Similar Work is required for Service/Works tenders")
    
    # Contract period validation
    if data['contract_period_months'] <= 0 and data['annualized_value'] <= 0:
        errors.append("Either Contract Period (months) or Annualized Estimated Value must be greater than 0")
    
    return len(errors) == 0, errors

def generate_bqc_document(data: Dict) -> Optional[BytesIO]:
    """Generate the BQC document in .docx format based on user inputs"""
    try:
        doc = Document()
        
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)
        
        # Calculate EMD amount
        emd_amount = calculate_emd(data['cec_estimate_excl_gst'], data['tender_type'])
        
        # Calculate turnover requirement (30% of CEC)
        turnover_requirement = 0.3 * data['cec_estimate_excl_gst']
        
        # Calculate experience requirements for Service/Works
        if data['tender_type'] in ['Service', 'Works']:
            option_a_value = 0.4 * data['cec_estimate_incl_gst']  # 40% of CEC (incl. GST)
            option_b_value = 0.5 * data['cec_estimate_incl_gst']  # 50% of CEC (incl. GST)
            option_c_value = 0.8 * data['cec_estimate_incl_gst']  # 80% of CEC (incl. GST)
        
        # Calculate Annualized Estimated Value
        if data['contract_period_months'] > 0:
            annualized_value = (data['cec_estimate_excl_gst'] / data['contract_period_months']) * 12
        else:
            annualized_value = data['annualized_value']
        
        # Header section with tabs for proper alignment
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.add_run("Ref: ").bold = True
        p.add_run(f"{data['ref_number']}")
        
        # Add tab for Date alignment
        tab_stop = p.paragraph_format.tab_stops.add_tab_stop(Inches(4.0), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)
        p.add_run("\tDate: ").bold = True
        p.add_run(f"{datetime.date.today().strftime('%d.%m.%Y')}")
        
        # NOTE TO line
        p = doc.add_paragraph()
        p.add_run("NOTE TO:\t").bold = True
        p.add_run(f"CHIEF PROCUREMENT OFFICER, CPO (M)/ PROCUREMENT LEADER GROUP {data['group_name']}")
        
        # SUBJECT line
        p = doc.add_paragraph()
        p.add_run("SUBJECT:\t").bold = True
        
        # Determine subject text based on tender type
        if data['tender_type'] == 'Goods':
            subject_text = f"SUPPLY OF '{data['item_name']}' FOR '{data['project_name']}' PROJECT"
        else:
            subject_text = f"JOB OF CONSTRUCTION OF '{data['project_name']}'"
        
        p.add_run(f"{subject_text}: APPROVAL OF BID QUALIFICATION CRITERIA AND FLOATING OF OPEN DOMESTIC TENDER.")
        
        # Add some space
        doc.add_paragraph("\t\t")
        
        # Section 1: PREAMBLE
        p = doc.add_paragraph()
        p.add_run("1.\tPREAMBLE").bold = True
        
        # Create table for PREAMBLE
        table = doc.add_table(rows=6, cols=2)
        table.style = 'Table Grid'
        
        # Fill table with data
        table.cell(0, 0).text = "Tender Description"
        table.cell(0, 1).text = data['tender_description']
        
        table.cell(1, 0).text = "PR reference/ Email reference"
        table.cell(1, 1).text = data['pr_reference']
        
        table.cell(2, 0).text = "Type of Tender"
        table.cell(2, 1).text = data['tender_type']
        
        table.cell(3, 0).text = "CEC estimate (incl. of GST)/ Date"
        table.cell(3, 1).text = f"{data['cec_estimate_incl_gst']} / {data['cec_date'].strftime('%d.%m.%Y')}"
        
        table.cell(4, 0).text = "CEC estimate exclusive of GST"
        table.cell(4, 1).text = str(data['cec_estimate_excl_gst'])
        
        table.cell(5, 0).text = "Budget Details (WBS/ Revex)"
        table.cell(5, 1).text = data['budget_details']
        
        p = doc.add_paragraph()
        p.add_run("Tender Platform – ").bold = True
        p.add_run(data['tender_platform'])
        
        # Add some space
        doc.add_paragraph("\t")
        
        # Section 2: BRIEF SCOPE OF WORK/ SUPPLY ITEMS
        p = doc.add_paragraph()
        p.add_run("2.\tBRIEF SCOPE OF WORK/ SUPPLY ITEMS").bold = True
        
        # Determine number of rows based on tender type
        if data['tender_type'] == 'Goods':
            rows = 6
        else:
            rows = 4
        
        # Create table for BRIEF SCOPE
        table = doc.add_table(rows=rows, cols=2)
        table.style = 'Table Grid'
        
        # Fill table with data
        table.cell(0, 0).text = "Brief Scope of Work / Supply Items"
        table.cell(0, 1).text = data['scope_of_work']
        
        table.cell(1, 0).text = "Contract Period /Completion Period"
        table.cell(1, 1).text = data['contract_period']
        
        if data['tender_type'] == 'Goods':
            table.cell(2, 0).text = "Delivery Period of the Item"
            table.cell(2, 1).text = data['delivery_period']
            
            table.cell(3, 0).text = "Warranty Period"
            table.cell(3, 1).text = data['warranty_period']
            
            table.cell(4, 0).text = "AMC/ CAMC/ O&M (No. of Years)"
            table.cell(4, 1).text = data['amc_period']
            
            table.cell(5, 0).text = "Payment Terms (if different from standard terms i.e within 30 days)"
            table.cell(5, 1).text = data['payment_terms']
        else:
            table.cell(2, 0).text = "AMC/ CAMC/ O&M (No. of Years)"
            table.cell(2, 1).text = data['amc_period']
            
            table.cell(3, 0).text = "Payment Terms (if different from standard terms i.e within 30 days)"
            table.cell(3, 1).text = data['payment_terms']
        
        # Add some space
        doc.add_paragraph("\t")
        
        # Section 3: BID QUALIFICATION CRITERIA (BQC)
        p = doc.add_paragraph()
        p.add_run("3.\tBID QUALIFICATION CRITERIA (BQC)").bold = True
        p = doc.add_paragraph("BPCL would like to qualify vendors for undertaking the above work as indicated in the brief scope. Detailed bid qualification criteria for short listing vendors shall be as follows:")
        
        # Technical Criteria based on tender type
        if data['tender_type'] == 'Goods':
            p = doc.add_paragraph()
            p.add_run("3.1\tTECHNICAL CRITERIA").bold = True
            p = doc.add_paragraph()
            p.add_run("3.1.1. For GOODS:").bold = True
            
            # Manufacturing Capability
            p = doc.add_paragraph()
            p.add_run("a) Manufacturing Capability:").bold = True
            p = doc.add_paragraph("Bidder* should be ")
            
            # Add selected manufacturer types
            manufacturer_types = ", ".join(data['manufacturer_types'])
            p.add_run(manufacturer_types).italic = True
            p.add_run(" of the item being tendered. ")
            
            p = doc.add_paragraph("*The definition of bidder is the entity which has a unique PAN (Permanent Account Number). All documents should be in the name of the bidder only (except in cases where the bidder is allowed to take the technical credentials of their OEM). Documents in the name of any legal entity other than the bidder, as defined above, shall not be accepted")
            
            p = doc.add_paragraph("(Explanatory Note: This is to ensure that the items are procured from the established Manufacturers, thereby ensuring right quality and price. OEMs are generally required to quote/participate in the tender. However, on case to case basis traders, authorized dealers, distributors, etc. may be considered depending on market scenario and volume of the tender.)")
            
            # Supplying Capacity
            p = doc.add_paragraph()
            p.add_run("b) Supplying Capacity:").bold = True
            
            # Calculate supplying capacity with MSE relaxation if applicable
            supplying_capacity = data['supplying_capacity']
            if data['mse_relaxation']:
                supplying_capacity = supplying_capacity * 0.85  # 15% relaxation
            
            p = doc.add_paragraph(f"The bidder shall have experience of having successfully supplied minimum of {supplying_capacity}% quantity of the annualized estimated quantity in any 12 continuous months during last 7 years in India or abroad, ending on last day of the month previous to the one in which tender is invited.")
            
            if data['mse_relaxation']:
                p = doc.add_paragraph("For MSE bidders Relaxation of 15% on the supplying capacity shall be given as per Corp. Finance Circular MA.TEC.POL.CON.3A dated 26.10.2020.")
            
            p = doc.add_paragraph("(Explanatory Note: The above criterion will demonstrate vendor's capability to manufacture the tendered item satisfactorily Based on market search 30% quantity of the annualized estimated quantity may be relaxed for widening competition for specific items)")
            
            p = doc.add_paragraph("*The definition of bidder is the entity which has a unique PAN (Permanent Account Number). All documents should be in the name of the bidder only (except in cases where the bidder is allowed to take the technical credentials of their OEM). Documents in the name of any legal entity other than the bidder, as defined above, shall not be accepted.")
            
        else:  # Service or Works
            p = doc.add_paragraph()
            p.add_run("3.1.2. BQC/PQC for Procurement of Works and Services:").bold = True
            
            p = doc.add_paragraph()
            p.add_run("I) Experience / Past performance / Technical Capability:").bold = True
            
            p = doc.add_paragraph("The bidder# should have experience of having successfully completed similar works during last 7 years ending last day of month previous to the one in which tender is floated should be either of the following: -")
            
            p = doc.add_paragraph(f"a. Three similar completed works each costing not less than the amount equal to Rs. {option_a_value:.2f} Lacs.")
            p = doc.add_paragraph("or")
            p = doc.add_paragraph(f"b. Two similar completed works each costing not less than the amount equal to Rs. {option_b_value:.2f} Lacs.")
            p = doc.add_paragraph("or")
            p = doc.add_paragraph(f"c. One similar completed work costing not less than the amount equal to Rs. {option_c_value:.2f} Lacs.")
            
            p = doc.add_paragraph(f'Definition of "similar work" should be clearly defined: {data["similar_work_definition"]}')
            
            p = doc.add_paragraph("# In case of Service contracts the term bidder may be suitably modified to take care of OEMs/ System Integrators/ Authorised Channel Partner etc.")
            p = doc.add_paragraph("*The definition of bidder is the entity which has a unique PAN (Permanent Account Number). All documents should be in the name of the bidder only (except in cases where the bidder is allowed to take the technical credentials of their OEM). Documents in the name of any legal entity other than the bidder, as defined above, shall not be accepted.")
            
            p = doc.add_paragraph("[ The % mentioned for meeting the similar works value can be amended considering the market inputs for specific items]")
            p = doc.add_paragraph("[In addition to above, additional criteria may be incorporated according to the requirement of the Project].")
        
        # Financial Criteria
        p = doc.add_paragraph()
        p.add_run("3.2\tFINANCIAL CRITERIA").bold = True
        
        p = doc.add_paragraph()
        p.add_run("3.2.1\tAVERAGE ANNUAL TURNOVER").bold = True
        p = doc.add_paragraph(f"The average annual turnover of the Bidder for last three audited accounting years shall be equal to or more than Rs. {turnover_requirement:.2f} Lacs.")
        
        p = doc.add_paragraph("Explanatory Notes:")
        p = doc.add_paragraph("i. Average annual turnover values in-line with CTE Office Memorandum No. 12-02-1-CTE-6 dated 17th Dec 2002.")
        
        p = doc.add_paragraph()
        p.add_run("3.2.2\tNET WORTH").bold = True
        p = doc.add_paragraph("The bidder should have positive net worth as per the latest audited financial statement.")
        
        p = doc.add_paragraph("Documents Required: Please refer the ITB (Instruction to Bidders) which mentions the documents to be submitted by bidders for meeting the above Technical and Financial criteria.")
        
        # BIDS MAY BE SUBMITTED BY
        p = doc.add_paragraph()
        p.add_run("3.3\tBIDS MAY BE SUBMITTED BY").bold = True
        
        p = doc.add_paragraph()
        p.add_run("3.3.1\t").bold = True
        p.add_run("An entity (domestic bidder) should have completed 3 financial years of existence as on original due date of tender since date of commencement of business and shall fulfil each BQC eligibility criteria as mentioned above.")
        
        p = doc.add_paragraph()
        p.add_run("3.3.2\t").bold = True
        p.add_run("JV/Consortium bids will not be accepted (i.e. Qualification on the strength of the JV Partners/Consortium Members /Subsidiaries / Group members will not be accepted)")
        
        # ESCALATION/ DE-ESCALATION CLAUSE
        if data['escalation_clause']:
            p = doc.add_paragraph()
            p.add_run("4. ESCALATION/ DE-ESCALATION CLAUSE: ").bold = True
            p.add_run(data['escalation_clause'])
        
        # EVALUATION METHODOLOGY
        p = doc.add_paragraph()
        p.add_run("5.\tEVALUATION METHODOLOGY").bold = True
        
        p = doc.add_paragraph("The tender will be invited through Open tender (Domestic) as two-part bid. The bid qualification evaluation of the received bids will be done as per the above bid qualification criteria and the technical bid of the shortlisted bidders will be evaluated subsequently. The price bids of the bidders who qualify BQC criteria & meet Technical / Commercial requirements of the tender will only be opened and evaluated.")
        
        p = doc.add_paragraph(f"The Commercial Evaluation shall be done on {data['evaluation_methodology']} basis. [Buyer to select the evaluation methodology]")
        
        p = doc.add_paragraph("The order will be placed based on above methodology AND Purchase preference based on MSE/ PPP-MII Policy.")
        
        p = doc.add_paragraph(f"The subject job is {data['divisibility']}. [for Purchase Preference for MSE and PPPMII; Buyer to select the option]")
        
        # EARNEST MONEY DEPOSIT (EMD)
        p = doc.add_paragraph()
        p.add_run("6.\tEARNEST MONEY DEPOSIT (EMD)").bold = True
        
        # Format EMD amount (0 for Nil)
        if emd_amount == 0:
            emd_text = "Nil"
        else:
            emd_text = f"{emd_amount} Lacs"
        
        p = doc.add_paragraph(f"Bidders are required to provide Earnest Money Deposit equivalent to Rs. {emd_text} for the tender.")
        
        p = doc.add_paragraph("EMD exemption shall be as per General Terms & Conditions of GeM (applicable for GeM tenders)/ MSE policy")
        
        p = doc.add_paragraph("Explanatory Note:")
        p = doc.add_paragraph("Procurement Group to justify the EMD amount as per Guidelines.")
        
        # Performance Security
        p = doc.add_paragraph()
        p.add_run("7. Performance Security (if at variance with the ITB clause):").bold = True
        
        # Determine standard performance security percentage
        if data['tender_type'] in ['Goods', 'Services']:
            standard_ps = 5
        else:  # Works
            standard_ps = 10
        
        if data['performance_security'] != standard_ps:
            p.add_run(f"Performance Security {data['performance_security']}% (approved by the competent authority).")
        else:
            p.add_run("Performance Security as per standard terms (5% for Goods & Services, 10% for Works).")
        
        # Other Points
        p = doc.add_paragraph("Other Points which may be taken into consideration while framing BQC:")
        p = doc.add_paragraph("1) Any guidelines from company, govt., industry tender practices (in case of industry tenders) etc. shall need to be followed superseding the above criteria as applicable.")
        p = doc.add_paragraph("2) Any services rendered by the vendor after due supply of the goods like AMC/CAMC after warranty period, servicing, etc. needs to be appropriately excluded while fixing the qualification (technical) criteria.")
        p = doc.add_paragraph("3) Where, the tender involves Annual Maintenance Contract (AMC) or Comprehensive Annual Maintenance Contract (CAMC), the estimated cost towards AMC/CAMC shall be excluded while arriving at the financial criteria (Annual Turnover) for the tender.")
        p = doc.add_paragraph("4) Additional qualification criteria may be built upon depending on the situation on case to case basis.")
        p = doc.add_paragraph("5) During first time procurement of any goods/services by CPO (M), inputs from User SBU/Entity should be taken. However same may be taken for subsequent procurements also to the extent possible and depending on the complexity of the procurement.")
        
        # APPROVAL REQUIRED
        p = doc.add_paragraph()
        p.add_run("7. APPROVAL REQUIRED").bold = True
        
        p = doc.add_paragraph("In view of above, approval is requested for the Supply of ")
        p.add_run(f"{data['item_name']}/ job -{data['project_name']}").bold = True
        p.add_run(" for:")
        
        p = doc.add_paragraph()
        p.add_run("i.\t").bold = True
        p.add_run("Bid Qualification Criteria as per Sr. No. 3, as per Clause 13.8 of Guidelines for procurement of Goods and Contract Services.")
        
        p = doc.add_paragraph()
        p.add_run("ii.\t").bold = True
        p.add_run("Inviting bids (two-part bid) through a Domestic Open Tender and adopting evaluation methodology as per Sr. No. 5 above.")
        
        p = doc.add_paragraph()
        p.add_run("iii.\t").bold = True
        p.add_run("Earnest Money Deposit as per Sr. No. 6 above./ Performance Security as per Sr. No. 7 (if applicable)")
        
        # Add some space
        doc.add_paragraph("\t")
        
        # Approval section
        p = doc.add_paragraph()
        p.add_run("Proposed by").bold = True
        p = doc.add_paragraph(f"{data['proposed_by']}, Procurement Manager (CPO Mktg.)")
        
        p = doc.add_paragraph()
        p.add_run("Recommended by").bold = True
        p = doc.add_paragraph(f"{data['recommended_by']}, Procurement Leader (CPO Mktg.)")
        
        p = doc.add_paragraph()
        p.add_run("Concurred by").bold = True
        p = doc.add_paragraph(f"{data['concurred_by']}, DGM Finance (CPO Mktg.)")
        
        p = doc.add_paragraph()
        p.add_run("Approved by").bold = True
        p = doc.add_paragraph(f"{data['approved_by']}, Chief Procurement Officer, (CPO Mktg.)")
        
        # Save the document to a bytes buffer
        doc_bytes = BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        
        logger.info(f"Successfully generated BQC document for ref: {data['ref_number']}")
        return doc_bytes
    
    except Exception as e:
        logger.error(f"Error generating BQC document: {str(e)}", exc_info=True)
        st.error(f"An error occurred while generating the document: {str(e)}")
        return None

def initialize_session_state():
    """Initialize session state with default values"""
    if 'data' not in st.session_state:
        st.session_state.data = {
            'ref_number': '',
            'group_name': '1 - Materials',
            'item_name': '',
            'project_name': '',
            'tender_description': '',
            'pr_reference': '',
            'tender_type': 'Goods',
            'cec_estimate_incl_gst': 0,
            'cec_date': datetime.date.today(),
            'cec_estimate_excl_gst': 0,
            'budget_details': '',
            'tender_platform': 'GeM',
            'scope_of_work': '',
            'contract_period': '',
            'contract_period_months': 0,
            'delivery_period': '',
            'warranty_period': '',
            'amc_period': '',
            'payment_terms': '',
            'manufacturer_types': ['Original Equipment Manufacturer'],
            'supplying_capacity': 30,
            'mse_relaxation': False,
            'similar_work_definition': '',
            'annualized_value': 0,
            'escalation_clause': '',
            'evaluation_methodology': 'Overall Lowest',
            'divisibility': 'Non-Divisible',
            'performance_security': 5,
            'proposed_by': 'XXXXX',
            'recommended_by': 'XXXXX',
            'concurred_by': 'Rajesh J.',
            'approved_by': 'Kani Amudhan N.'
        }
    
    if 'document_generated' not in st.session_state:
        st.session_state.document_generated = False
    
    if 'form_errors' not in st.session_state:
        st.session_state.form_errors = []

def clear_form():
    """Clear all form fields"""
    initialize_session_state()
    st.session_state.document_generated = False
    st.session_state.form_errors = []
    st.success("Form cleared successfully!")
    logger.info("Form cleared by user")

def main():
    st.set_page_config(
        page_title="BQC Document Generator",
        page_icon="📝",
        layout="wide",
        initial_sidebar_state="expanded"
    )
    
    st.title("BQC Document Generator")
    st.write("Automated generation of Bid Qualification Criteria documents")
    
    # Initialize session state
    initialize_session_state()
    
    # Display form errors if any
    if st.session_state.form_errors:
        st.error("Please fix the following errors:")
        for error in st.session_state.form_errors:
            st.error(f"• {error}")
    
    # Create form for data input
    with st.form("bqc_form"):
        st.header("PREAMBLE")
        col1, col2 = st.columns(2)
        with col1:
            st.session_state.data['ref_number'] = st.text_input(
                "Ref Number", 
                value=st.session_state.data['ref_number'], 
                placeholder="XXXXXX",
                help="Unique reference number for this BQC"
            )
            
            # Group dropdown with names
            selected_group = st.selectbox(
                "Procurement Group", 
                options=list(GROUP_OPTIONS.keys()),
                format_func=lambda x: f"{x} - {GROUP_OPTIONS[x]}",
                index=0
            )
            st.session_state.data['group_name'] = f"{selected_group} - {GROUP_OPTIONS[selected_group]}"
            
            st.session_state.data['item_name'] = st.text_input(
                "Item Name", 
                value=st.session_state.data['item_name'], 
                placeholder="ABC",
                help="Name of the item/service being procured"
            )
            
            # Dynamic label for project name based on tender type
            if st.session_state.data['tender_type'] == 'Goods':
                project_label = "Good to be procured"
                project_placeholder = "XXX"
            else:
                project_label = "Service to be given"
                project_placeholder = "XXX"
                
            st.session_state.data['project_name'] = st.text_input(
                project_label, 
                value=st.session_state.data['project_name'], 
                placeholder=project_placeholder
            )
            
            st.session_state.data['tender_description'] = st.text_area(
                "Tender Description", 
                value=st.session_state.data['tender_description'],
                help="Detailed description of the tender"
            )
            st.session_state.data['pr_reference'] = st.text_input(
                "PR Reference/Email Reference", 
                value=st.session_state.data['pr_reference'],
                help="Purchase requisition reference or email reference"
            )
        with col2:
            # Tender type selection with state management
            tender_type = st.selectbox(
                "Type of Tender", 
                TENDER_TYPES, 
                index=0,
                help="Select the type of tender"
            )
            if tender_type != st.session_state.data['tender_type']:
                st.session_state.data['tender_type'] = tender_type
                st.rerun()  # This will refresh the app when tender type changes
            
            st.session_state.data['cec_estimate_incl_gst'] = st.number_input(
                "CEC Estimate (incl. GST) (in Lakh)", 
                min_value=0, 
                value=st.session_state.data['cec_estimate_incl_gst'],
                help="Cost estimate including GST"
            )
            st.session_state.data['cec_date'] = st.date_input(
                "CEC Date", 
                value=st.session_state.data['cec_date'],
                help="Date of cost estimate"
            )
            st.session_state.data['cec_estimate_excl_gst'] = st.number_input(
                "CEC Estimate (excl. GST) (in Lakh)", 
                min_value=0, 
                value=st.session_state.data['cec_estimate_excl_gst'],
                help="Cost estimate excluding GST"
            )
            st.session_state.data['budget_details'] = st.text_input(
                "Budget Details (WBS/Revex)", 
                value=st.session_state.data['budget_details'],
                help="Work breakdown structure or revenue expenditure details"
            )
            st.session_state.data['tender_platform'] = st.selectbox(
                "Tender Platform", 
                PLATFORM_OPTIONS, 
                index=0,
                help="Platform where tender will be floated"
            )
        
        st.header("BRIEF SCOPE OF WORK/ SUPPLY ITEMS")
        col1, col2 = st.columns(2)
        with col1:
            st.session_state.data['scope_of_work'] = st.text_area(
                "Brief Scope of Work/Supply Items", 
                value=st.session_state.data['scope_of_work'],
                help="Brief description of work or items to be supplied"
            )
            st.session_state.data['contract_period'] = st.text_input(
                "Contract Period/Completion Period", 
                value=st.session_state.data['contract_period'],
                help="Duration of the contract or time for completion"
            )
            st.session_state.data['contract_period_months'] = st.number_input(
                "Contract Period (in months)", 
                min_value=0, 
                value=st.session_state.data['contract_period_months'],
                help="Duration of contract in months (for annualized calculation)"
            )
            st.session_state.data['amc_period'] = st.text_input(
                "AMC/CAMC/O&M (No. of Years)", 
                value=st.session_state.data['amc_period'],
                help="Annual maintenance contract period in years"
            )
            st.session_state.data['payment_terms'] = st.text_input(
                "Payment Terms (if different from standard)", 
                value=st.session_state.data['payment_terms'],
                help="Payment terms if different from standard 30 days"
            )
        with col2:
            if st.session_state.data['tender_type'] == 'Goods':
                st.session_state.data['delivery_period'] = st.text_input(
                    "Delivery Period of the Item", 
                    value=st.session_state.data['delivery_period'],
                    help="Time required for delivery of items"
                )
                st.session_state.data['warranty_period'] = st.text_input(
                    "Warranty Period", 
                    value=st.session_state.data['warranty_period'],
                    help="Warranty period for the supplied items"
                )
        
        # Calculate and display Annualized Estimated Value
        if st.session_state.data['contract_period_months'] > 0:
            annualized_value = (st.session_state.data['cec_estimate_excl_gst'] / st.session_state.data['contract_period_months']) * 12
            st.session_state.data['annualized_value'] = annualized_value
            st.info(f"Annualized Estimated Value: Rs. {annualized_value:.2f} Lacs")
        else:
            st.session_state.data['annualized_value'] = st.number_input(
                "Annualized Estimated Value (in Lakh)", 
                min_value=0, 
                value=st.session_state.data['annualized_value'],
                help="Used to calculate turnover requirement (30%)"
            )
        
        st.header("BID QUALIFICATION CRITERIA (BQC)")
        
        # Technical Criteria based on tender type
        if st.session_state.data['tender_type'] == 'Goods':
            st.subheader("Technical Criteria for Goods")
            st.session_state.data['manufacturer_types'] = st.multiselect(
                "Manufacturer Types",
                MANUFACTURER_TYPES,
                default=st.session_state.data['manufacturer_types'],
                help="Select applicable manufacturer types"
            )
            st.session_state.data['supplying_capacity'] = st.number_input(
                "Supplying Capacity (%)", 
                min_value=0, 
                max_value=100, 
                value=st.session_state.data['supplying_capacity'],
                help="Percentage of annualized estimated quantity"
            )
            st.session_state.data['mse_relaxation'] = st.checkbox(
                "Apply MSE Relaxation (15%)", 
                value=st.session_state.data['mse_relaxation'],
                help="Check to apply 15% relaxation for MSE bidders"
            )
            
            # Show MSE relaxation calculation
            if st.session_state.data['mse_relaxation']:
                relaxed_capacity = st.session_state.data['supplying_capacity'] * 0.85
                st.info(f"Relaxed Supplying Capacity (15% relaxation): {relaxed_capacity}%")
        else:
            st.subheader("Technical Criteria for Service/Works")
            # Calculate and display the experience requirements
            option_a_value = 0.4 * st.session_state.data['cec_estimate_incl_gst']  # 40% of CEC (incl. GST)
            option_b_value = 0.5 * st.session_state.data['cec_estimate_incl_gst']  # 50% of CEC (incl. GST)
            option_c_value = 0.8 * st.session_state.data['cec_estimate_incl_gst']  # 80% of CEC (incl. GST)
            
            st.info(f"Option A (40%): Rs. {option_a_value:.2f} Lacs")
            st.info(f"Option B (50%): Rs. {option_b_value:.2f} Lacs")
            st.info(f"Option C (80%): Rs. {option_c_value:.2f} Lacs")
            
            st.session_state.data['similar_work_definition'] = st.text_area(
                "Definition of Similar Work", 
                value=st.session_state.data['similar_work_definition'],
                help="Define what constitutes 'similar work' for this tender"
            )
        
        # Financial Criteria
        st.subheader("Financial Criteria")
        
        # Calculate and display 30% of CEC
        turnover_requirement = 0.3 * st.session_state.data['cec_estimate_excl_gst']
        st.info(f"Annual Turnover Requirement (30% of CEC): Rs. {turnover_requirement:.2f} Lacs")
        
        st.header("OTHER SECTIONS")
        col1, col2 = st.columns(2)
        with col1:
            st.session_state.data['escalation_clause'] = st.text_area(
                "Escalation/De-escalation Clause (if applicable)", 
                value=st.session_state.data['escalation_clause'],
                help="Details of any price escalation or de-escalation clause"
            )
            
            # Add explanation for Escalation/De-escalation Clause
            # with st.expander("What is an Escalation/De-escalation Clause?"):
            #     st.markdown("""
            #     **Escalation Clause**: Allows price increase if costs (materials, labor, etc.) rise significantly during contract period.
                
            #     **De-escalation Clause**: Allows price decrease if costs fall significantly during contract period.
                
            #     **Why use it?**
            #     - For long-term contracts with volatile market conditions
            #     - Protects suppliers from unexpected cost increases
            #     - Protects buyers from overpaying when costs decrease
            #     - Makes contracts fairer for both parties
                
            #     **How it works:**
            #     - Specifies which costs can be adjusted
            #     - Defines how price changes are calculated
            #     - Sets thresholds for when adjustments apply
            #     - May limit maximum price changes
            #     """)
            
            st.session_state.data['evaluation_methodology'] = st.selectbox(
                "Evaluation Methodology",
                EVALUATION_METHODS,
                index=0,
                help="Method to be used for evaluating bids"
            )
            st.session_state.data['divisibility'] = st.selectbox(
                "Divisibility", 
                DIVISIBILITY_OPTIONS, 
                index=0,
                help="Whether the tender can be divided into parts"
            )
        with col2:
            # Show EMD calculation preview
            emd_preview = calculate_emd(
                st.session_state.data['cec_estimate_excl_gst'], 
                st.session_state.data['tender_type']
            )
            if emd_preview == 0:
                emd_preview_text = "Nil"
            else:
                emd_preview_text = f"{emd_preview} Lacs"
            
            st.info(f"EMD will be automatically calculated as: Rs. {emd_preview_text}")
            
            # Set default performance security based on tender type
            if st.session_state.data['tender_type'] in ['Goods', 'Services']:
                default_ps = 5
            else:  # Works
                default_ps = 10
                
            st.session_state.data['performance_security'] = st.number_input(
                "Performance Security (%)", 
                min_value=0, 
                max_value=20, 
                value=default_ps,
                help="5% for Goods & Services, 10% for Works (standard)"
            )
        
        st.header("APPROVAL SECTION")
        col1, col2 = st.columns(2)
        with col1:
            st.session_state.data['proposed_by'] = st.text_input(
                "Proposed By", 
                value=st.session_state.data['proposed_by'],
                help="Name of person proposing the BQC"
            )
            st.session_state.data['recommended_by'] = st.text_input(
                "Recommended By", 
                value=st.session_state.data['recommended_by'],
                help="Name of person recommending the BQC"
            )
        with col2:
            st.session_state.data['concurred_by'] = st.text_input(
                "Concurred By", 
                value=st.session_state.data['concurred_by'],
                help="Name of person concurring the BQC"
            )
            st.session_state.data['approved_by'] = st.text_input(
                "Approved By", 
                value=st.session_state.data['approved_by'],
                help="Name of person approving the BQC"
            )
        
        # Add submit button
        submitted = st.form_submit_button("Generate Document")
    
    # Add Clear All Form button at the end
    if st.button("Clear All Form"):
        clear_form()
    
    # Process form submission
    if submitted:
        # Validate form data
        is_valid, errors = validate_input(st.session_state.data)
        st.session_state.form_errors = errors
        
        if is_valid:
            # Generate the document
            doc_bytes = generate_bqc_document(st.session_state.data)
            
            if doc_bytes:
                # Set document generated flag
                st.session_state.document_generated = True
                
                # Display success message only after clicking Generate
                st.success("Document generated successfully!")
                
                # Provide download button
                st.download_button(
                    label="Download Document",
                    data=doc_bytes,
                    file_name=f"BQC_{st.session_state.data['ref_number']}_{datetime.date.today().strftime('%Y%m%d')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        else:
            st.error("Please fix the errors above before generating the document")

if __name__ == "__main__":
    main()
